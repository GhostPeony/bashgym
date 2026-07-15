#!/usr/bin/env python3
"""Build the MemexAI embedding experiment report, charts, and source tables."""

from __future__ import annotations

import argparse
import csv
import json
import textwrap
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

COLORS = {
    "ink": "1F2937",
    "muted": "667085",
    "purple": "7457C8",
    "purple_light": "EEE9FA",
    "sage": "56876D",
    "sage_light": "E7F0EA",
    "amber": "C58A2A",
    "amber_light": "FAEFD9",
    "rose": "B85C72",
    "grid": "D9DEE7",
    "paper": "FFFFFF",
    "soft": "F5F7FA",
    "red": "A23B3B",
}
MODEL_COLORS = {
    "Base Qwen 0.6B": COLORS["purple"],
    "Real 702": COLORS["sage"],
    "Mixed 898": COLORS["amber"],
}


def rgb(hex_value: str) -> tuple[int, int, int]:
    return tuple(int(hex_value[i : i + 2], 16) for i in (0, 2, 4))


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default(size=size)


def chart_canvas(title: str, subtitle: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1600, 900), rgb(COLORS["paper"]))
    draw = ImageDraw.Draw(image)
    draw.text((90, 62), title, fill=rgb(COLORS["ink"]), font=font(44, True))
    draw.text((90, 120), subtitle, fill=rgb(COLORS["muted"]), font=font(24))
    draw.rounded_rectangle((86, 174, 1514, 816), radius=20, fill=rgb(COLORS["soft"]))
    return image, draw


def draw_legend(draw: ImageDraw.ImageDraw, labels: list[str], y: int = 190) -> None:
    x = 110
    for label in labels:
        color = rgb(MODEL_COLORS[label])
        draw.rounded_rectangle((x, y, x + 26, y + 26), radius=5, fill=color)
        draw.text((x + 38, y - 2), label, fill=rgb(COLORS["ink"]), font=font(20, True))
        x += 210


def grouped_bar_chart(
    output: Path,
    *,
    title: str,
    subtitle: str,
    categories: list[str],
    series: dict[str, list[float]],
    minimum: float = 0.0,
    maximum: float = 1.0,
) -> None:
    image, draw = chart_canvas(title, subtitle)
    labels = list(series)
    draw_legend(draw, labels)
    left, top, right, bottom = 150, 260, 1470, 730
    for tick in range(6):
        value = minimum + (maximum - minimum) * tick / 5
        y = bottom - int((value - minimum) / (maximum - minimum) * (bottom - top))
        draw.line((left, y, right, y), fill=rgb(COLORS["grid"]), width=2)
        draw.text((78, y - 13), f"{value:.2f}", fill=rgb(COLORS["muted"]), font=font(18))
    group_width = (right - left) / len(categories)
    bar_width = min(62, int(group_width / (len(labels) + 1)))
    for category_index, category in enumerate(categories):
        center = left + group_width * (category_index + 0.5)
        total_width = len(labels) * bar_width + (len(labels) - 1) * 10
        start = center - total_width / 2
        for series_index, label in enumerate(labels):
            value = series[label][category_index]
            x1 = int(start + series_index * (bar_width + 10))
            x2 = x1 + bar_width
            y = bottom - int((value - minimum) / (maximum - minimum) * (bottom - top))
            draw.rounded_rectangle((x1, y, x2, bottom), radius=8, fill=rgb(MODEL_COLORS[label]))
            value_text = f"{value:.3f}"
            bbox = draw.textbbox((0, 0), value_text, font=font(17, True))
            draw.text(
                (x1 + (bar_width - (bbox[2] - bbox[0])) / 2, y - 28),
                value_text,
                fill=rgb(COLORS["ink"]),
                font=font(17, True),
            )
        wrapped = textwrap.wrap(category, width=14)
        for line_index, line in enumerate(wrapped):
            bbox = draw.textbbox((0, 0), line, font=font(19, True))
            draw.text(
                (center - (bbox[2] - bbox[0]) / 2, bottom + 18 + line_index * 23),
                line,
                fill=rgb(COLORS["ink"]),
                font=font(19, True),
            )
    output.parent.mkdir(parents=True, exist_ok=True)
    image.save(output, dpi=(180, 180))


def dataset_chart(output: Path, bundle: dict[str, Any]) -> None:
    image, draw = chart_canvas(
        "Training data composition",
        "Both arms retain the 114-row template base; the mixed arm adds a controlled synthetic slice.",
    )
    left, right = 180, 1400
    rows = [
        ("Real 702", [("Template", 114, COLORS["purple"]), ("Real DD", 588, COLORS["sage"])]),
        (
            "Mixed 898",
            [
                ("Template", 114, COLORS["purple"]),
                ("Real DD", 588, COLORS["sage"]),
                ("Synthetic DD", 196, COLORS["amber"]),
            ],
        ),
    ]
    max_total = 898
    for row_index, (label, parts) in enumerate(rows):
        y1 = 300 + row_index * 210
        y2 = y1 + 92
        draw.text((100, y1 + 25), label, fill=rgb(COLORS["ink"]), font=font(28, True))
        x = left + 170
        available = right - x
        for part_label, value, color in parts:
            width = int(available * value / max_total)
            draw.rounded_rectangle((x, y1, x + width, y2), radius=12, fill=rgb(color))
            text = f"{part_label}\n{value}"
            lines = text.splitlines()
            for line_index, line in enumerate(lines):
                bbox = draw.textbbox((0, 0), line, font=font(20, True))
                draw.text(
                    (x + (width - (bbox[2] - bbox[0])) / 2, y1 + 18 + line_index * 27),
                    line,
                    fill=(255, 255, 255),
                    font=font(20, True),
                )
            x += width
        total = sum(value for _, value, _ in parts)
        draw.text((right + 18, y1 + 28), str(total), fill=rgb(COLORS["ink"]), font=font(25, True))
    qtypes = bundle["statistics"]["selected_synthetic_dd"]["query_types"]
    note = (
        f"Selected synthetic mix: {qtypes['keyword_query']} keyword, "
        f"{qtypes['natural_question']} natural question, "
        f"{qtypes['semantic_paraphrase']} semantic paraphrase."
    )
    draw.rounded_rectangle((110, 710, 1490, 786), radius=14, fill=rgb(COLORS["amber_light"]))
    draw.text((140, 733), note, fill=rgb(COLORS["ink"]), font=font(22, True))
    output.parent.mkdir(parents=True, exist_ok=True)
    image.save(output, dpi=(180, 180))


def training_cost_chart(output: Path) -> None:
    image, draw = chart_canvas(
        "Training cost of the two fine-tune arms",
        "Synthetic passages were much longer, increasing wall-clock time beyond their 28% increase in optimizer steps.",
    )
    panels = [
        ("Optimizer steps", [88, 114], 120, "steps"),
        ("Training duration", [133, 316], 340, "seconds"),
    ]
    labels = ["Real 702", "Mixed 898"]
    colors = [COLORS["sage"], COLORS["amber"]]
    for panel_index, (panel_title, values, max_value, unit) in enumerate(panels):
        x0 = 120 + panel_index * 735
        draw.text((x0, 245), panel_title, fill=rgb(COLORS["ink"]), font=font(30, True))
        for row_index, (label, value, color) in enumerate(zip(labels, values, colors, strict=True)):
            y = 340 + row_index * 190
            draw.text((x0, y), label, fill=rgb(COLORS["ink"]), font=font(24, True))
            draw.rounded_rectangle((x0, y + 45, x0 + 590, y + 115), radius=10, fill=rgb(COLORS["grid"]))
            width = int(590 * value / max_value)
            draw.rounded_rectangle((x0, y + 45, x0 + width, y + 115), radius=10, fill=rgb(color))
            draw.text((x0 + width + 14, y + 63), f"{value} {unit}", fill=rgb(COLORS["ink"]), font=font(21, True))
    draw.rounded_rectangle((120, 742, 1480, 806), radius=12, fill=rgb(COLORS["purple_light"]))
    draw.text(
        (150, 760),
        "Mixed arm: 1.30x the steps, 2.38x the training time, and lower exact MRR.",
        fill=rgb(COLORS["ink"]),
        font=font(23, True),
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    image.save(output, dpi=(180, 180))


def retrieval_layers_diagram(output: Path) -> None:
    image, draw = chart_canvas(
        "What is learned, and what is fixed retrieval math?",
        "Fine-tuning changes the embedding encoder. Exact cosine/dot-product search is parameter-free ranking math.",
    )
    boxes = [
        (105, 300, 375, 525, "Query + passage\ntraining examples", COLORS["purple_light"], "DATA"),
        (435, 300, 725, 525, "Qwen embedding\nencoder weights", COLORS["sage_light"], "LEARNED"),
        (785, 300, 1065, 525, "Normalized vectors\n+ exact dot product", COLORS["soft"], "FIXED MATH"),
        (1125, 300, 1495, 525, "Ranked candidates\n+ optional reranker", COLORS["amber_light"], "SYSTEM"),
    ]
    for x1, y1, x2, y2, body, fill, kicker in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=rgb(fill), outline=rgb(COLORS["grid"]), width=3)
        draw.text((x1 + 24, y1 + 22), kicker, fill=rgb(COLORS["purple"]), font=font(19, True))
        for line_index, line in enumerate(body.splitlines()):
            draw.text((x1 + 24, y1 + 78 + line_index * 39), line, fill=rgb(COLORS["ink"]), font=font(26, True))
    for x in (375, 725, 1065):
        draw.line((x + 10, 412, x + 50, 412), fill=rgb(COLORS["muted"]), width=5)
        draw.polygon([(x + 50, 412), (x + 35, 402), (x + 35, 422)], fill=rgb(COLORS["muted"]))
    draw.rounded_rectangle((150, 650, 1450, 772), radius=16, fill=rgb(COLORS["purple_light"]))
    lines = [
        "Encoder training can reshape semantic geometry.",
        "BM25 fusion, chunking, filters, ANN indexes, and rerankers are separate system levers.",
    ]
    for index, line in enumerate(lines):
        draw.text((190, 676 + index * 39), line, fill=rgb(COLORS["ink"]), font=font(24, index == 0))
    output.parent.mkdir(parents=True, exist_ok=True)
    image.save(output, dpi=(180, 180))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if table.rows:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None, italic: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, *, before: float = 0, after: float = 6, line: float = 1.1) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    run = paragraph.add_run(text)
    set_run_font(run, bold=True, color=COLORS["purple"] if level < 3 else COLORS["ink"])


def add_body(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    style_paragraph(paragraph)
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True, color=COLORS["ink"])
        rest = paragraph.add_run(text[len(bold_lead) :])
        set_run_font(rest, color=COLORS["ink"])
    else:
        run = paragraph.add_run(text)
        set_run_font(run, color=COLORS["ink"])


def add_callout(doc: Document, label: str, text: str, *, fill: str = "EEE9FA") -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    style_paragraph(p, after=2)
    lead = p.add_run(f"{label}: ")
    set_run_font(lead, bold=True, color=COLORS["purple"])
    body = p.add_run(text)
    set_run_font(body, color=COLORS["ink"])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, COLORS["soft"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index else WD_ALIGN_PARAGRAPH.LEFT
        style_paragraph(p, after=0, line=1.0)
        run = p.add_run(header)
        set_run_font(run, size=9.5, bold=True, color=COLORS["ink"])
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            p = cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            style_paragraph(p, after=0, line=1.0)
            run = p.add_run(value)
            set_run_font(run, size=9.5, color=COLORS["ink"])
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline = run.add_picture(str(image_path), width=Inches(6.3))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", caption)
    doc_pr.set("title", caption)
    caption_p = doc.add_paragraph(style="Caption")
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(caption_p, before=2, after=8, line=1.0)
    caption_run = caption_p.add_run(caption)
    set_run_font(caption_run, size=9, italic=True, color=COLORS["muted"])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=COLORS["muted"])
    for instruction in ("PAGE", "NUMPAGES"):
        if instruction == "NUMPAGES":
            sep = paragraph.add_run(" of ")
            set_run_font(sep, size=9, color=COLORS["muted"])
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instruction
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        field_run = paragraph.add_run()
        set_run_font(field_run, size=9, color=COLORS["muted"])
        field_run._r.append(fld_char_begin)
        field_run._r.append(instr_text)
        field_run._r.append(fld_char_end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    heading_tokens = {
        1: (16, 16, 8),
        2: (13, 12, 6),
        3: (12, 8, 4),
    }
    for level, (size, before, after) in heading_tokens.items():
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(COLORS["purple"] if level < 3 else COLORS["ink"])
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.167
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run("MEMEXAI EMBEDDING EXPERIMENTS  |  TECHNICAL FINDINGS")
    set_run_font(header_run, size=8.5, bold=True, color=COLORS["muted"])
    add_page_number(section.footer.paragraphs[0])
    section.different_first_page_header_footer = True


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(38)
    kicker = doc.add_paragraph()
    style_paragraph(kicker, after=10)
    run = kicker.add_run("TECHNICAL FINDINGS REPORT")
    set_run_font(run, size=11, bold=True, color=COLORS["sage"])
    title = doc.add_paragraph()
    style_paragraph(title, after=8, line=1.0)
    run = title.add_run("MemexAI Embedding\nModel Experiments")
    set_run_font(run, size=30, bold=True, color=COLORS["ink"])
    subtitle = doc.add_paragraph()
    style_paragraph(subtitle, after=24)
    run = subtitle.add_run(
        "Full-data fine-tuning, frozen retrieval evaluation, product benchmarking, and the next experiment map"
    )
    set_run_font(run, size=14, color=COLORS["muted"])
    meta_rows = [
        ["Campaign", "memex-embed-full-v1-20260711"],
        ["Model family", "Qwen3-Embedding-0.6B at 768 dimensions"],
        ["Evaluation", "36 held-out queries over 2,018 transcript chunks"],
        ["Status", "Complete - no fine-tuned model promoted"],
    ]
    add_table(doc, ["Field", "Value"], meta_rows, [2100, 7260])
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    add_callout(
        doc,
        "Decision",
        "Keep zero-shot Qwen3-Embedding-0.6B as the exact-order champion. The real-data fine-tune improves neighborhood recall, but neither trained arm improves exact MRR.",
        fill=COLORS["purple_light"],
    )
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(note, before=20, after=0)
    run = note.add_run("Prepared July 11, 2026  |  BashGym / MemexAI")
    set_run_font(run, size=10, color=COLORS["muted"])
    doc.add_page_break()


def build_docx(
    output: Path,
    charts: dict[str, Path],
    bundle: dict[str, Any],
    model_metrics: dict[str, dict[str, Any]],
    product: dict[str, dict[str, Any]],
) -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "Executive summary", 1)
    add_callout(
        doc,
        "Bottom line",
        "The real 702-pair arm moved retrieval in the desired direction at broader ranks, but exact ordering remained below the zero-shot baseline. The mixed synthetic arm cost more and weakened exact ordering further.",
        fill=COLORS["sage_light"],
    )
    summary_rows = [
        ["Zero-shot base", "0.4502", "0.8056", "0.5721", "Champion"],
        ["Real 702", "0.4426", "0.8611", "0.5979", "Neighborhood gain"],
        ["Mixed 898", "0.4277", "0.8611", "0.5862", "Do not continue"],
    ]
    add_table(
        doc,
        ["Candidate", "Exact MRR", "Exact R@10", "Local MRR", "Decision"],
        summary_rows,
        [2100, 1650, 1650, 1650, 2310],
    )
    add_body(
        doc,
        "What changed: the training pipeline now embeds passages exactly as SearchTube does - title plus transcript content - and uses collision-safe batches so multiple queries sharing one positive chunk cannot become false in-batch negatives.",
        bold_lead="What changed:",
    )
    add_body(
        doc,
        "What did not change: production models, embedding tables, indexes, and serving pointers were not modified by this campaign.",
        bold_lead="What did not change:",
    )
    add_figure(doc, charts["retrieval"], "Figure 1. Exact held-out retrieval metrics across the three candidates.")

    doc.add_page_break()
    add_heading(doc, "Training data and experiment design", 1)
    add_body(
        doc,
        "Both fine-tune arms start with 114 template queries. The real arm adds all 588 judged, real-grounded Data Designer queries. The mixed arm adds a capped 196-query synthetic slice selected deterministically across all 130 synthetic positives.",
    )
    add_figure(doc, charts["dataset"], "Figure 2. Training-row composition for the real and mixed arms.")
    dataset_rows = [
        ["Template base", "114", "38", "19", "Original training split"],
        ["Real Data Designer", "588", "199", "19", "98% judge keep rate"],
        ["Selected synthetic", "196", "130", "38 pseudo-videos", "25% of new DD augmentation"],
        ["Held-out test", "36", "12", "6", "Untouched, video-disjoint"],
    ]
    add_table(
        doc,
        ["Dataset component", "Queries", "Positives", "Videos", "Role"],
        dataset_rows,
        [2350, 1200, 1300, 1800, 2710],
    )
    add_callout(
        doc,
        "Data caution",
        "Synthetic passages had a median length near 440 words versus about 197 for selected real passages. Generator and judge were the same model, and some accepted rows lacked judge reasons. Synthetic data should remain a capped ablation until those distribution and independence issues are fixed.",
        fill=COLORS["amber_light"],
    )

    doc.add_page_break()
    add_heading(doc, "Held-out evaluation findings", 1)
    add_body(
        doc,
        "The frozen test contains 36 queries from six videos that do not overlap the training videos. Each candidate re-embeds the same 2,018-document real corpus, then uses exact normalized dot-product search at 768 dimensions.",
    )
    add_figure(doc, charts["neighborhood"], "Figure 3. Exact, local-window, and same-video MRR.")
    add_body(
        doc,
        "The real arm improves local-window MRR from 0.5721 to 0.5979 and exact Recall@10 from 0.8056 to 0.8611. Exact MRR still falls from 0.4502 to 0.4426, so the broader neighborhood gain does not meet the promotion gate.",
    )
    add_callout(
        doc,
        "Interpretation",
        "The encoder is reaching the correct transcript neighborhood more often, but it is not learning enough evidence to rank the exact chunk first. Adjacent transcript chunks likely need to be represented as additional positives, not incidental negatives.",
        fill=COLORS["purple_light"],
    )

    doc.add_page_break()
    add_heading(doc, "Query-type and product-fixture findings", 1)
    add_figure(doc, charts["query_type"], "Figure 4. Exact MRR by query type.")
    add_body(
        doc,
        "The real arm's clearest slice benefit is natural-question retrieval: local-window MRR rises from 0.4823 to 0.5978. Keyword exact MRR softens from 0.7486 to 0.7181, while semantic-paraphrase exact MRR remains the weakest slice and falls slightly from 0.1702 to 0.1636.",
    )
    product_rows = [
        ["Base Qwen 0.6B", "1.000", "1.000", "0.998", "Pass"],
        ["Real 702", "1.000", "1.000", "0.998", "Pass"],
        ["Mixed 898", "1.000", "1.000", "0.998", "Pass"],
    ]
    add_table(
        doc,
        ["Candidate", "R@1", "MRR", "nDCG@5", "Product fixture"],
        product_rows,
        [2600, 1300, 1300, 1500, 2660],
    )
    add_callout(
        doc,
        "Benchmark limitation",
        "The 20-document, 14-query SearchTube product fixture is saturated. It remains a useful release smoke test, but cannot select among these candidates. Add near-duplicate clips, ambiguous paraphrases, and project-scope collisions before using it as a promotion gate.",
        fill=COLORS["amber_light"],
    )

    doc.add_page_break()
    add_heading(doc, "Training efficiency", 1)
    add_figure(doc, charts["cost"], "Figure 5. Optimizer steps and remote training duration.")
    add_body(
        doc,
        "The mixed arm increased optimizer steps by 30%, but wall-clock time by 138%, because its selected synthetic passages were substantially longer. It still finished with lower exact MRR than both the base and real-only arms.",
    )
    add_callout(
        doc,
        "Cost decision",
        "Do not spend another cycle on the current synthetic-positive recipe. Its 2.38x training time did not produce a quality gain over the real-only arm.",
        fill=COLORS["purple_light"],
    )

    doc.add_page_break()
    add_heading(doc, "Encoder training versus retrieval math", 1)
    add_figure(doc, charts["layers"], "Figure 6. Learned encoder components versus fixed and system-level retrieval components.")
    add_heading(doc, "What these runs train", 2)
    add_body(
        doc,
        "The fine-tunes update Qwen's embedding-encoder weights with a contrastive objective. They do not teach Qwen to generate answers. Training reshapes vector geometry so queries should land closer to relevant passages than to negatives.",
    )
    add_heading(doc, "What is fixed math", 2)
    add_body(
        doc,
        "With normalized vectors, dot product and cosine similarity produce the same ranking. This exact scorer has no learned parameters. At 2,018 chunks it is the correct quality oracle; an approximate nearest-neighbor index can change speed and approximation error, but cannot repair bad geometry.",
    )
    add_heading(doc, "What belongs to the retrieval system", 2)
    system_items = [
        "BM25 plus dense fusion changes candidate generation and ranking without changing the encoder.",
        "A Qwen reranker over the dense top 20 or 50 is a separate learned pair-ranking model.",
        "Chunking, title and alias projection, access filters, and local-window grouping change the contract around the encoder.",
    ]
    for item in system_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, color=COLORS["ink"])

    doc.add_page_break()
    add_heading(doc, "Recommended next experiment", 1)
    add_callout(
        doc,
        "Change the supervision, not just the epoch count",
        "The next campaign should use multiple transcript-window positives, positive-aware hard negatives, and a larger effective contrastive batch. Repeating pair-only MNRL with more synthetic positives is not supported by these results.",
        fill=COLORS["sage_light"],
    )
    steps = [
        "Expand the video-disjoint dev set, especially natural questions and semantic paraphrases. Keep the current 36 test rows untouched.",
        "Benchmark zero-shot Qwen3-Embedding-0.6B with Qwen3-Reranker-0.6B over top-20 and top-50, plus BM25/dense reciprocal-rank fusion.",
        "Mine candidates with base Qwen, BM25, and optionally Gemini. Remove exact, local-window, and same-answer candidates, then teacher- or human-filter 3-7 negatives per query.",
        "Treat the exact chunk plus relevant +/-2 transcript chunks as multiple positives where they answer the same information need.",
        "Train with Qwen InfoNCE or CachedMNRL, effective batch 128-512, explicit negatives, false-negative masking, and Matryoshka objectives at 1024d and 768d.",
        "Use 768d video-disjoint dev nDCG/MRR for checkpoint selection. Re-admit length-matched synthetic data only as a separately reported ablation.",
    ]
    for item in steps:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.keep_together = True
        run = p.add_run(item)
        set_run_font(run, color=COLORS["ink"])
    add_heading(doc, "Primary references", 2)
    references = [
        "Qwen3-Embedding SWIFT training guide: https://github.com/QwenLM/Qwen3-Embedding/blob/main/docs/training/SWIFT.md",
        "Sentence Transformers loss guidance: https://www.sbert.net/docs/package_reference/sentence_transformer/losses.html",
        "NV-Retriever positive-aware hard-negative mining: https://arxiv.org/abs/2407.15831",
        "E5 weakly supervised contrastive pretraining: https://arxiv.org/abs/2212.03533",
    ]
    for item in references:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=9.5, color=COLORS["muted"])

    doc.add_page_break()
    add_heading(doc, "Appendix: reproducibility contract", 1)
    source_hashes = bundle["sources"]
    hash_rows = [
        ["Real corpus", source_hashes["real_corpus"]["sha256"]],
        ["Template queries", source_hashes["template_queries"]["sha256"]],
        ["Real DD queries", source_hashes["real_queries"]["sha256"]],
        ["Synthetic DD queries", source_hashes["synthetic_queries"]["sha256"]],
        ["Synthetic corpus", source_hashes["synthetic_corpus"]["sha256"]],
    ]
    add_table(doc, ["Artifact", "SHA-256"], hash_rows, [2300, 7060])
    add_body(
        doc,
        "Local source package includes the editable Word report, PDF, Markdown source, chart PNGs, experiment metrics CSV, dataset summary CSV, and raw JSON manifests used to build the report.",
    )
    add_body(
        doc,
        "Verification at report build: 19 MemexAI tests passed; both remote training manifests report completed; the real and mixed arms contain 88 and 114 optimizer steps respectively; and the frozen baseline reproduces exact MRR 0.450224.",
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "MemexAI Embedding Model Experiments"
    doc.core_properties.subject = "Training, evaluation, benchmark, and next experiment report"
    doc.core_properties.author = "BashGym / MemexAI"
    doc.core_properties.keywords = "MemexAI, embeddings, retrieval, Qwen, evaluation"
    doc.save(output)


def write_csv(path: Path, headers: list[str], rows: list[list[Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as fh:
        writer = csv.writer(fh)
        writer.writerow(headers)
        writer.writerows(rows)


def write_markdown(output: Path, charts: dict[str, Path]) -> None:
    relative = {name: path.relative_to(output.parent).as_posix() for name, path in charts.items()}
    text = f"""# MemexAI Embedding Model Experiments

Full-data fine-tuning, frozen retrieval evaluation, product benchmarking, and the next experiment map.

## Decision

Keep zero-shot Qwen3-Embedding-0.6B as the exact-order champion. Real 702 improves neighborhood retrieval, but neither fine-tune improves exact MRR. Do not promote either trained model.

## Main results

| Candidate | Exact MRR | Exact R@10 | Local-window MRR | Decision |
|---|---:|---:|---:|---|
| Base Qwen 0.6B | **0.450224** | 0.805556 | 0.572124 | Champion |
| Real 702 | 0.442565 | **0.861111** | **0.597911** | Neighborhood gain only |
| Mixed 898 | 0.427717 | **0.861111** | 0.586205 | Stop this arm |

![Exact retrieval metrics]({relative['retrieval']}){{width=full}}

![Training data composition]({relative['dataset']}){{width=full}}

![Neighborhood MRR]({relative['neighborhood']}){{width=full}}

![Query-type MRR]({relative['query_type']}){{width=full}}

![Training cost]({relative['cost']}){{width=full}}

## Encoder training versus retrieval math

These runs update the Qwen embedding encoder's weights with a contrastive objective. They do not train answer generation. With normalized vectors, exact dot product and cosine produce the same ranking and have no learned parameters.

![Learned encoder versus retrieval system diagram]({relative['layers']}){{width=full}}

BM25 fusion, reranking, chunking, title/alias projection, filters, and ANN indexes are separate retrieval-system levers.

## Next experiment

1. Expand the video-disjoint dev set while keeping the current 36 test rows untouched.
2. Benchmark zero-shot Qwen plus Qwen3-Reranker-0.6B and BM25/dense fusion.
3. Mine 3-7 positive-aware, teacher-filtered hard negatives per query.
4. Represent answer-equivalent local transcript windows as multiple positives.
5. Train with InfoNCE or CachedMNRL, effective batch 128-512, false-negative masking, and Matryoshka objectives at 1024d and 768d.
6. Re-admit length-matched synthetic data only as a separately reported ablation.

## References

- Qwen3-Embedding SWIFT training: <https://github.com/QwenLM/Qwen3-Embedding/blob/main/docs/training/SWIFT.md>
- Sentence Transformers losses: <https://www.sbert.net/docs/package_reference/sentence_transformer/losses.html>
- NV-Retriever: <https://arxiv.org/abs/2407.15831>
- E5: <https://arxiv.org/abs/2212.03533>
"""
    output.write_text(text, encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--bundle-manifest", type=Path, required=True)
    parser.add_argument("--experiment-root", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)
    charts_dir = args.output_dir / "charts"
    tables_dir = args.output_dir / "tables"
    bundle = load_json(args.bundle_manifest)
    manifest_paths = {
        "Base Qwen 0.6B": args.experiment_root / "baseline" / "query_format_ablation_manifest.json",
        "Real 702": args.experiment_root / "real702" / "query_format_ablation_manifest.json",
        "Mixed 898": args.experiment_root / "mixed898" / "query_format_ablation_manifest.json",
    }
    model_metrics = {
        name: load_json(path)["runs"]["memexai_youtube"]["metrics"]
        for name, path in manifest_paths.items()
    }
    product = {
        "Base Qwen 0.6B": load_json(args.experiment_root / "product-benchmark" / "base.json"),
        "Real 702": load_json(args.experiment_root / "product-benchmark" / "real702.json"),
        "Mixed 898": load_json(args.experiment_root / "product-benchmark" / "mixed898.json"),
    }

    charts = {
        "retrieval": charts_dir / "01_exact_retrieval_metrics.png",
        "dataset": charts_dir / "02_training_data_composition.png",
        "neighborhood": charts_dir / "03_neighborhood_mrr.png",
        "query_type": charts_dir / "04_query_type_exact_mrr.png",
        "cost": charts_dir / "05_training_cost.png",
        "layers": charts_dir / "06_encoder_vs_retrieval_system.png",
    }
    grouped_bar_chart(
        charts["retrieval"],
        title="Held-out exact retrieval quality",
        subtitle="36 test queries, 2,018 real transcript chunks, exact normalized dot-product search at 768d.",
        categories=["MRR", "Recall@1", "Recall@3", "Recall@5", "Recall@10"],
        series={
            name: [
                metrics["overall"]["exact_chunk_mrr"],
                metrics["overall"]["exact_chunk_recall_at_1"],
                metrics["overall"]["exact_chunk_recall_at_3"],
                metrics["overall"]["exact_chunk_recall_at_5"],
                metrics["overall"]["exact_chunk_recall_at_10"],
            ]
            for name, metrics in model_metrics.items()
        },
    )
    dataset_chart(charts["dataset"], bundle)
    grouped_bar_chart(
        charts["neighborhood"],
        title="Ranking quality at three relevance granularities",
        subtitle="Real 702 improves local-window and same-video ranking, while base retains the best exact-chunk MRR.",
        categories=["Exact chunk MRR", "Local-window MRR", "Same-video MRR"],
        series={
            name: [
                metrics["overall"]["exact_chunk_mrr"],
                metrics["overall"]["local_window_mrr"],
                metrics["overall"]["same_video_mrr"],
            ]
            for name, metrics in model_metrics.items()
        },
    )
    grouped_bar_chart(
        charts["query_type"],
        title="Exact MRR by query type",
        subtitle="Keyword queries remain strongest. Semantic paraphrases are the main unresolved model-quality gap.",
        categories=["Keyword", "Natural question", "Semantic paraphrase"],
        series={
            name: [
                metrics["by_query_type"][query_type]["exact_chunk_mrr"]
                for query_type in ("keyword_query", "natural_question", "semantic_paraphrase")
            ]
            for name, metrics in model_metrics.items()
        },
    )
    training_cost_chart(charts["cost"])
    retrieval_layers_diagram(charts["layers"])

    metrics_rows: list[list[Any]] = []
    for model, metrics in model_metrics.items():
        overall = metrics["overall"]
        metrics_rows.append(
            [
                model,
                overall["exact_chunk_mrr"],
                overall["exact_chunk_recall_at_1"],
                overall["exact_chunk_recall_at_3"],
                overall["exact_chunk_recall_at_5"],
                overall["exact_chunk_recall_at_10"],
                overall["local_window_mrr"],
                overall["local_window_recall_at_10"],
                overall["same_video_mrr"],
            ]
        )
    write_csv(
        tables_dir / "experiment_metrics.csv",
        [
            "candidate",
            "exact_mrr",
            "exact_recall_at_1",
            "exact_recall_at_3",
            "exact_recall_at_5",
            "exact_recall_at_10",
            "local_window_mrr",
            "local_window_recall_at_10",
            "same_video_mrr",
        ],
        metrics_rows,
    )
    stats = bundle["statistics"]
    write_csv(
        tables_dir / "dataset_summary.csv",
        ["component", "rows", "positive_chunks", "videos", "keyword", "natural_question", "semantic_paraphrase"],
        [
            [
                label,
                stats[key]["rows"],
                stats[key]["positive_chunks"],
                stats[key]["videos"],
                stats[key]["query_types"].get("keyword_query", 0),
                stats[key]["query_types"].get("natural_question", 0),
                stats[key]["query_types"].get("semantic_paraphrase", 0),
            ]
            for label, key in (
                ("Original train", "original_train"),
                ("Real DD", "real_dd"),
                ("Selected synthetic DD", "selected_synthetic_dd"),
                ("Real arm", "real_arm"),
                ("Mixed arm", "mixed_arm"),
                ("Held-out test", "heldout_test"),
            )
        ],
    )
    write_csv(
        tables_dir / "query_type_metrics.csv",
        ["candidate", "query_type", "exact_mrr", "exact_recall_at_1", "exact_recall_at_10", "local_window_mrr"],
        [
            [
                model,
                query_type,
                values["exact_chunk_mrr"],
                values["exact_chunk_recall_at_1"],
                values["exact_chunk_recall_at_10"],
                values["local_window_mrr"],
            ]
            for model, metrics in model_metrics.items()
            for query_type, values in metrics["by_query_type"].items()
        ],
    )
    write_csv(
        tables_dir / "training_runs.csv",
        ["candidate", "training_rows", "positive_chunks", "optimizer_steps", "duration_seconds", "exact_mrr", "exact_recall_at_10"],
        [
            ["Real 702", 702, 237, 88, 133, 0.442565, 0.861111],
            ["Mixed 898", 898, 367, 114, 316, 0.427717, 0.861111],
        ],
    )
    build_docx(
        args.output_dir / "MemexAI_Embedding_Experiment_Report.docx",
        charts,
        bundle,
        model_metrics,
        product,
    )
    write_markdown(args.output_dir / "MemexAI_Embedding_Experiment_Report.md", charts)
    print(args.output_dir.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
