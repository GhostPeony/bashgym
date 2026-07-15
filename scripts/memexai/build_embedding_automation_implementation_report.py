#!/usr/bin/env python3
"""Build the evidence-backed MemexAI embedding automation implementation report.

The report intentionally separates the complete Candidate B quality experiment from
the one-step Ponyo campaign smoke. The former can support retrieval findings; the
latter proves remote campaign ownership, restart adoption, sealing, and accounting.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import html
import json
import statistics
import zipfile
from collections.abc import Iterable, Sequence
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

REPORT_STEM = "MemexAI_Embedding_Automation_Implementation_Report"
REPORT_TITLE = "MemexAI Embedding Automation Campaign"
REPORT_SUBTITLE = (
    "Candidate B training and retrieval findings, durable Ponyo orchestration proof, "
    "milestone audit, and next-iteration plan"
)

COLORS = {
    "paper": "FCFBF7",
    "ink": "1F2937",
    "muted": "667085",
    "purple": "6D5BD0",
    "purple_light": "EEE9FA",
    "sage": "6F9D78",
    "sage_light": "E6F1E8",
    "amber": "B7791F",
    "amber_light": "FFF4D6",
    "rose": "B95C67",
    "soft": "F2F4F7",
    "grid": "D9DEE7",
}


@dataclass(frozen=True)
class ReportInputs:
    candidate_run_dir: Path
    base_dev_manifest: Path
    candidate_a_dev_manifest: Path
    campaign_evidence: Path
    campaign_export_manifest: Path


def load_json(path: Path) -> dict[str, Any]:
    value = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(value, dict):
        raise ValueError(f"expected a JSON object: {path}")
    return value


def load_jsonl(path: Path) -> list[dict[str, Any]]:
    rows = [
        json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()
    ]
    if not all(isinstance(row, dict) for row in rows):
        raise ValueError(f"expected JSON objects in: {path}")
    return rows


def load_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def write_csv(path: Path, headers: Sequence[str], rows: Iterable[Sequence[Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle, lineterminator="\n")
        writer.writerow(headers)
        writer.writerows(rows)


def rolling_mean(values: Sequence[float], window: int = 10) -> list[float]:
    if window < 1:
        raise ValueError("rolling window must be >= 1")
    return [
        statistics.fmean(values[max(0, index - window + 1) : index + 1])
        for index in range(len(values))
    ]


def summarize_training_metrics(
    manifest: dict[str, Any], rows: list[dict[str, Any]]
) -> dict[str, Any]:
    loss_rows = sorted(
        (row for row in rows if "loss" in row and "step" in row),
        key=lambda row: int(row["step"]),
    )
    expected_steps = int(manifest["training"]["optimizer_steps"])
    steps = [int(row["step"]) for row in loss_rows]
    if steps != list(range(1, expected_steps + 1)):
        raise ValueError(
            f"full-run loss series must contain steps 1..{expected_steps}; got {steps[:1]}..{steps[-1:]}"
        )
    losses = [float(row["loss"]) for row in loss_rows]
    if expected_steps <= 1:
        raise ValueError("Candidate B quality evidence must be a full run, not a smoke")
    return {
        "steps": steps,
        "epoch_progress": [float(row["epoch"]) for row in loss_rows],
        "losses": losses,
        "rolling_loss": rolling_mean(losses, 10),
        "learning_rates": [float(row["learning_rate"]) for row in loss_rows],
        "grad_norms": [float(row["grad_norm"]) for row in loss_rows],
        "first_loss": losses[0],
        "last_loss": losses[-1],
        "minimum_loss": min(losses),
        "mean_logged_loss": statistics.fmean(losses),
        "epoch_boundary_step": int(manifest["training"]["batches_per_epoch"]),
    }


def summarize_resources(rows: list[dict[str, str]]) -> dict[str, Any]:
    if not rows:
        raise ValueError("resource sample CSV is empty")
    times = [float(row["epoch_utc"]) for row in rows]
    available = [float(row["mem_available_bytes"]) / (1024**3) for row in rows]
    swap = [float(row["swap_used_bytes"]) / (1024**3) for row in rows]
    gpu = [float(row["gpu_util_pct"]) for row in rows]
    temperatures = [float(row["temp_c"]) for row in rows]
    start = times[0]
    return {
        "minutes": [(value - start) / 60 for value in times],
        "available_memory_gib": available,
        "swap_used_gib": swap,
        "gpu_utilization_pct": gpu,
        "temperature_c": temperatures,
        "sample_count": len(rows),
        "available_memory_floor_gib": min(available),
        "swap_growth_gib": max(swap) - swap[0],
        "gpu_utilization_mean_pct": statistics.fmean(gpu),
        "temperature_peak_c": max(temperatures),
    }


def strict_dev_metrics(manifest: dict[str, Any]) -> dict[str, Any]:
    if manifest.get("rows") != 18 or manifest.get("splits") != ["dev"]:
        raise ValueError("strict development evidence must be the 18-row dev-only split")
    query_path = str(manifest.get("queries_jsonl", ""))
    if Path(query_path).name != "heldout-dev.jsonl":
        raise ValueError("strict development evidence must use heldout-dev.jsonl")
    run = manifest["runs"]["memexai_youtube"]
    metrics = dict(run["metrics"]["overall"])
    metrics["median_query_latency_ms"] = float(run["median_query_latency_ms"])
    metrics["model_footprint_bytes"] = int(manifest["model_footprint_bytes"])
    return metrics


def verify_export_manifest(path: Path, manifest: dict[str, Any]) -> list[dict[str, Any]]:
    verified: list[dict[str, Any]] = []
    for item in manifest.get("files", []):
        target = path.parent / str(item["name"])
        if not target.is_file():
            raise ValueError(f"campaign export file is missing: {target}")
        size = target.stat().st_size
        digest = sha256_file(target)
        if size != int(item["size_bytes"]) or digest != str(item["sha256"]):
            raise ValueError(f"campaign export hash/size mismatch: {target}")
        verified.append({"name": target.name, "size_bytes": size, "sha256": digest})
    if not verified:
        raise ValueError("campaign export manifest lists no files")
    return verified


def milestone_rows(facts: dict[str, Any]) -> list[dict[str, str]]:
    campaign = facts["campaign"]
    export = facts["campaign_export"]
    rows = [
        {
            "milestone": "Candidate B full cached-MNRL training",
            "status": "complete",
            "evidence": f"{facts['training']['optimizer_steps']} optimizer steps; completed manifest",
        },
        {
            "milestone": "Protected dev-only retrieval comparison",
            "status": "complete",
            "evidence": "Base, Candidate A, and Candidate B each use heldout-dev.jsonl (18 rows)",
        },
        {
            "milestone": "Product retrieval compatibility fixture",
            "status": "complete",
            "evidence": "Candidate B fixture passed; treated as saturated compatibility evidence",
        },
        {
            "milestone": "Server-owned Ponyo campaign smoke",
            "status": "complete",
            "evidence": f"attempt {campaign['attempt_id']} completed with one remote launch",
        },
        {
            "milestone": "Worker-restart adoption of exact remote identity",
            "status": "complete",
            "evidence": f"{campaign['remote_adoption_event_count']} adoption events; claim generation {campaign['claim_generation']}",
        },
        {
            "milestone": "Sealed outputs, budget accounting, and no automatic promotion",
            "status": "complete",
            "evidence": f"{campaign['sealed_artifact_count']} artifacts; {campaign['budget']['actual']:.2f} GPU-hours; promotion=false",
        },
        {
            "milestone": "Campaign multi-format export integrity",
            "status": "complete",
            "evidence": f"{len(export['verified_files'])} files verified against SHA-256 manifest",
        },
        {
            "milestone": "Automated implementation report package",
            "status": "complete",
            "evidence": "Markdown, DOCX, PDF, CSV, PNG, SVG, and source/output manifest",
        },
        {
            "milestone": "BM25/dense RRF and independent reranker benchmark",
            "status": "not evidenced",
            "evidence": "No scored BM25, RRF, or reranker artifact supplied to this report",
        },
        {
            "milestone": "30-query double-label pilot",
            "status": "not evidenced",
            "evidence": "No annotation or agreement artifact supplied to this report",
        },
        {
            "milestone": "Persistent resident campaign worker service",
            "status": "not evidenced",
            "evidence": "Smoke harness proves replacement/adoption, not an installed long-running service",
        },
        {
            "milestone": "Same-campaign API, CLI/MCP, and live canvas proof",
            "status": "not evidenced",
            "evidence": "No integrated cross-surface trace or live Electron capture supplied",
        },
    ]
    return rows


def load_report_facts(inputs: ReportInputs) -> dict[str, Any]:
    candidate_manifest_path = inputs.candidate_run_dir / "training_manifest.json"
    candidate_metrics_path = inputs.candidate_run_dir / "training_metrics.jsonl"
    resource_path = inputs.candidate_run_dir / "resource_samples.csv"
    candidate_dev_path = (
        inputs.candidate_run_dir
        / "eval"
        / "dev"
        / "candidate-b"
        / "query_format_ablation_manifest.json"
    )
    product_path = inputs.candidate_run_dir / "eval" / "product" / "candidate-b.json"
    source_paths = [
        candidate_manifest_path,
        candidate_metrics_path,
        resource_path,
        inputs.base_dev_manifest,
        inputs.candidate_a_dev_manifest,
        candidate_dev_path,
        product_path,
        inputs.campaign_evidence,
        inputs.campaign_export_manifest,
    ]
    missing = [path for path in source_paths if not path.is_file()]
    if missing:
        raise FileNotFoundError("missing report evidence: " + ", ".join(map(str, missing)))

    training_manifest = load_json(candidate_manifest_path)
    if training_manifest.get("status") != "completed":
        raise ValueError("Candidate B training manifest is not completed")
    training_summary = summarize_training_metrics(
        training_manifest, load_jsonl(candidate_metrics_path)
    )
    resources = summarize_resources(load_csv(resource_path))
    base = strict_dev_metrics(load_json(inputs.base_dev_manifest))
    candidate_a = strict_dev_metrics(load_json(inputs.candidate_a_dev_manifest))
    candidate_b = strict_dev_metrics(load_json(candidate_dev_path))
    product = load_json(product_path)
    if not product.get("passed"):
        raise ValueError("Candidate B product fixture did not pass")
    campaign = load_json(inputs.campaign_evidence)
    required_campaign_truths = {
        "attempt_status": "completed",
        "single_remote_launch_event": True,
        "promotion_performed": False,
        "protected_test_opened": False,
    }
    for key, expected in required_campaign_truths.items():
        if campaign.get(key) != expected:
            raise ValueError(f"campaign evidence {key!r} must equal {expected!r}")
    loss_points = campaign.get("loss_points", [])
    if len(loss_points) != 1 or int(loss_points[0]["step"]) != 1:
        raise ValueError("campaign proof must remain a one-step orchestration smoke")
    export_manifest = load_json(inputs.campaign_export_manifest)
    verified_export_files = verify_export_manifest(inputs.campaign_export_manifest, export_manifest)
    source_hashes = [
        {
            "name": path.name,
            "role": _source_role(path, inputs),
            "path": str(path.resolve()),
            "sha256": sha256_file(path),
            "size_bytes": path.stat().st_size,
        }
        for path in source_paths
    ]
    facts: dict[str, Any] = {
        "training_manifest": training_manifest,
        "training": {**training_manifest["training"], **training_summary},
        "resources": resources,
        "models": {
            "Base Qwen 0.6B": base,
            "Candidate A": candidate_a,
            "Candidate B": candidate_b,
        },
        "product": product,
        "campaign": campaign,
        "campaign_export": {
            **export_manifest,
            "verified_files": verified_export_files,
        },
        "source_hashes": source_hashes,
        "evidence_cutoff": max(str(training_manifest["completed_at"]), str(campaign["created_at"])),
    }
    facts["milestones"] = milestone_rows(facts)
    return facts


def _source_role(path: Path, inputs: ReportInputs) -> str:
    if path == inputs.base_dev_manifest:
        return "strict_dev_base"
    if path == inputs.candidate_a_dev_manifest:
        return "strict_dev_candidate_a"
    if path == inputs.campaign_evidence:
        return "campaign_smoke_evidence"
    if path == inputs.campaign_export_manifest:
        return "campaign_export_manifest"
    if path.name == "training_manifest.json":
        return "candidate_b_training_manifest"
    if path.name == "training_metrics.jsonl":
        return "candidate_b_training_metrics"
    if path.name == "resource_samples.csv":
        return "candidate_b_resource_samples"
    if path.name == "candidate-b.json":
        return "candidate_b_product_fixture"
    return "strict_dev_candidate_b"


def _rgb(value: str) -> tuple[int, int, int]:
    return tuple(int(value[index : index + 2], 16) for index in (0, 2, 4))


def _font(size: int, bold: bool = False):
    from PIL import ImageFont

    choices = (
        [r"C:\Windows\Fonts\arialbd.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"]
        if bold
        else [r"C:\Windows\Fonts\arial.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"]
    )
    for choice in choices:
        if Path(choice).exists():
            return ImageFont.truetype(choice, size)
    return ImageFont.load_default()


def _canvas(title: str, subtitle: str):
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (1600, 900), _rgb(COLORS["paper"]))
    draw = ImageDraw.Draw(image)
    draw.text((88, 58), title, fill=_rgb(COLORS["ink"]), font=_font(42, True))
    draw.text((88, 116), subtitle, fill=_rgb(COLORS["muted"]), font=_font(23))
    draw.rounded_rectangle((84, 172, 1516, 818), radius=20, fill=_rgb(COLORS["soft"]))
    return image, draw


def _svg_text(
    x: float, y: float, text: str, size: int, *, bold: bool = False, color: str = "ink"
) -> str:
    weight = "700" if bold else "400"
    return (
        f'<text x="{x}" y="{y}" font-family="Arial, sans-serif" '
        f'font-size="{size}" font-weight="{weight}" fill="#{COLORS[color]}">'
        f"{html.escape(text)}</text>"
    )


def _write_svg(path: Path, body: Sequence[str], title: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    content = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<svg xmlns="http://www.w3.org/2000/svg" width="1600" height="900" viewBox="0 0 1600 900">',
        f"<title>{html.escape(title)}</title>",
        f'<rect width="1600" height="900" fill="#{COLORS["paper"]}"/>',
        *body,
        "</svg>",
    ]
    path.write_text("\n".join(content) + "\n", encoding="utf-8")


def build_loss_chart(output_dir: Path, facts: dict[str, Any]) -> tuple[Path, Path]:
    png = output_dir / "01_candidate_b_full_run_loss.png"
    svg = output_dir / "01_candidate_b_full_run_loss.svg"
    steps = facts["training"]["steps"]
    raw = facts["training"]["losses"]
    rolling = facts["training"]["rolling_loss"]
    left, top, right, bottom = 170, 280, 1460, 720
    maximum = max(raw) * 1.08
    image, draw = _canvas(
        "Candidate B training loss — complete 84-step run",
        "Raw cached-MNRL loss and 10-step rolling mean. The one-step campaign smoke is excluded.",
    )
    svg_body = [
        _svg_text(88, 92, "Candidate B training loss — complete 84-step run", 42, bold=True),
        _svg_text(
            88,
            140,
            "Raw cached-MNRL loss and 10-step rolling mean. The one-step campaign smoke is excluded.",
            23,
            color="muted",
        ),
        f'<rect x="84" y="172" width="1432" height="646" rx="20" fill="#{COLORS["soft"]}"/>',
    ]
    for tick in range(6):
        value = maximum * tick / 5
        y = bottom - (bottom - top) * tick / 5
        draw.line((left, y, right, y), fill=_rgb(COLORS["grid"]), width=2)
        draw.text((100, y - 12), f"{value:.2f}", fill=_rgb(COLORS["muted"]), font=_font(17))
        svg_body.append(
            f'<line x1="{left}" y1="{y:.2f}" x2="{right}" y2="{y:.2f}" stroke="#{COLORS["grid"]}" stroke-width="2"/>'
        )
        svg_body.append(_svg_text(100, y + 5, f"{value:.2f}", 17, color="muted"))

    def points(values: Sequence[float]) -> list[tuple[float, float]]:
        return [
            (
                left + (right - left) * (step - 1) / (steps[-1] - 1),
                bottom - (bottom - top) * value / maximum,
            )
            for step, value in zip(steps, values, strict=True)
        ]

    raw_points, rolling_points = points(raw), points(rolling)
    draw.line(raw_points, fill=_rgb(COLORS["rose"]), width=3)
    draw.line(rolling_points, fill=_rgb(COLORS["sage"]), width=6)
    raw_svg_points = " ".join(f"{x:.2f},{y:.2f}" for x, y in raw_points)
    rolling_svg_points = " ".join(f"{x:.2f},{y:.2f}" for x, y in rolling_points)
    svg_body.append(
        f'<polyline fill="none" stroke="#{COLORS["rose"]}" stroke-width="3" points="{raw_svg_points}"/>'
    )
    svg_body.append(
        f'<polyline fill="none" stroke="#{COLORS["sage"]}" stroke-width="6" points="{rolling_svg_points}"/>'
    )
    boundary = facts["training"]["epoch_boundary_step"]
    boundary_x = left + (right - left) * (boundary - 1) / (steps[-1] - 1)
    draw.line((boundary_x, top, boundary_x, bottom), fill=_rgb(COLORS["amber"]), width=3)
    draw.text(
        (boundary_x + 9, top + 9), "epoch 2", fill=_rgb(COLORS["amber"]), font=_font(17, True)
    )
    svg_body.extend(
        [
            f'<line x1="{boundary_x:.2f}" y1="{top}" x2="{boundary_x:.2f}" y2="{bottom}" stroke="#{COLORS["amber"]}" stroke-width="3"/>',
            _svg_text(boundary_x + 9, top + 27, "epoch 2", 17, bold=True, color="amber"),
        ]
    )
    for index, (label, color) in enumerate((("raw loss", "rose"), ("10-step mean", "sage"))):
        x = 490 + index * 330
        draw.line((x, 230, x + 55, 230), fill=_rgb(COLORS[color]), width=6)
        draw.text((x + 68, 216), label, fill=_rgb(COLORS["ink"]), font=_font(19, True))
        svg_body.extend(
            [
                f'<line x1="{x}" y1="230" x2="{x + 55}" y2="230" stroke="#{COLORS[color]}" stroke-width="6"/>',
                _svg_text(x + 68, 237, label, 19, bold=True),
            ]
        )
    draw.text((740, 750), "optimizer step", fill=_rgb(COLORS["muted"]), font=_font(18, True))
    svg_body.append(_svg_text(740, 772, "optimizer step", 18, bold=True, color="muted"))
    png.parent.mkdir(parents=True, exist_ok=True)
    image.save(png, format="PNG", dpi=(180, 180), optimize=False)
    _write_svg(svg, svg_body, "Candidate B full-run training loss")
    return png, svg


def build_dev_chart(output_dir: Path, facts: dict[str, Any]) -> tuple[Path, Path]:
    png = output_dir / "02_strict_dev_comparison.png"
    svg = output_dir / "02_strict_dev_comparison.svg"
    labels = ["Exact MRR", "Local-window MRR", "Same-video MRR"]
    keys = ["exact_chunk_mrr", "local_window_mrr", "same_video_mrr"]
    series = list(facts["models"].items())
    colors = ["purple", "amber", "sage"]
    left, top, right, bottom = 170, 300, 1460, 720
    image, draw = _canvas(
        "Strict held-out development retrieval",
        "18 dev-only queries over the same 2,018-chunk corpus. Product fixture results are not mixed into this chart.",
    )
    svg_body = [
        _svg_text(88, 92, "Strict held-out development retrieval", 42, bold=True),
        _svg_text(
            88,
            140,
            "18 dev-only queries over the same 2,018-chunk corpus. Product fixture results are not mixed into this chart.",
            23,
            color="muted",
        ),
        f'<rect x="84" y="172" width="1432" height="646" rx="20" fill="#{COLORS["soft"]}"/>',
    ]
    for tick in range(6):
        value = tick / 5
        y = bottom - value * (bottom - top)
        draw.line((left, y, right, y), fill=_rgb(COLORS["grid"]), width=2)
        draw.text((104, y - 11), f"{value:.1f}", fill=_rgb(COLORS["muted"]), font=_font(17))
        svg_body.extend(
            [
                f'<line x1="{left}" y1="{y:.2f}" x2="{right}" y2="{y:.2f}" stroke="#{COLORS["grid"]}" stroke-width="2"/>',
                _svg_text(104, y + 5, f"{value:.1f}", 17, color="muted"),
            ]
        )
    group_width = (right - left) / len(labels)
    for group_index, (label, key) in enumerate(zip(labels, keys, strict=True)):
        center = left + group_width * (group_index + 0.5)
        for series_index, ((name, metrics), color) in enumerate(zip(series, colors, strict=True)):
            value = float(metrics[key])
            x1 = center - 114 + series_index * 78
            x2 = x1 + 62
            y = bottom - value * (bottom - top)
            draw.rounded_rectangle((x1, y, x2, bottom), radius=7, fill=_rgb(COLORS[color]))
            draw.text(
                (x1 - 2, y - 25), f"{value:.3f}", fill=_rgb(COLORS["ink"]), font=_font(15, True)
            )
            svg_body.extend(
                [
                    f'<rect x="{x1:.2f}" y="{y:.2f}" width="62" height="{bottom - y:.2f}" rx="7" fill="#{COLORS[color]}"/>',
                    _svg_text(x1 - 2, y - 7, f"{value:.3f}", 15, bold=True),
                ]
            )
        draw.text((center - 90, bottom + 22), label, fill=_rgb(COLORS["ink"]), font=_font(18, True))
        svg_body.append(_svg_text(center - 90, bottom + 43, label, 18, bold=True))
    for index, ((name, _), color) in enumerate(zip(series, colors, strict=True)):
        x = 300 + index * 350
        draw.rounded_rectangle((x, 220, x + 26, 246), radius=4, fill=_rgb(COLORS[color]))
        draw.text((x + 38, 218), name, fill=_rgb(COLORS["ink"]), font=_font(18, True))
        svg_body.extend(
            [
                f'<rect x="{x}" y="220" width="26" height="26" rx="4" fill="#{COLORS[color]}"/>',
                _svg_text(x + 38, 239, name, 18, bold=True),
            ]
        )
    image.save(png, format="PNG", dpi=(180, 180), optimize=False)
    _write_svg(svg, svg_body, "Strict development comparison")
    return png, svg


def build_resource_chart(output_dir: Path, facts: dict[str, Any]) -> tuple[Path, Path]:
    png = output_dir / "03_candidate_b_resource_trace.png"
    svg = output_dir / "03_candidate_b_resource_trace.svg"
    resources = facts["resources"]
    minutes = resources["minutes"]
    panels = [
        (
            "Available unified memory (GiB)",
            resources["available_memory_gib"],
            "purple",
            max(resources["available_memory_gib"]) * 1.05,
            (170, 285, 1460, 480),
        ),
        (
            "GPU utilization (%)",
            resources["gpu_utilization_pct"],
            "sage",
            100.0,
            (170, 590, 1460, 755),
        ),
    ]
    image, draw = _canvas(
        "Candidate B Ponyo resource trace",
        f"{resources['sample_count']} samples across the full training run; panels use independent labeled scales.",
    )
    svg_body = [
        _svg_text(88, 92, "Candidate B Ponyo resource trace", 42, bold=True),
        _svg_text(
            88,
            140,
            f"{resources['sample_count']} samples across the full training run; panels use independent labeled scales.",
            23,
            color="muted",
        ),
        f'<rect x="84" y="172" width="1432" height="646" rx="20" fill="#{COLORS["soft"]}"/>',
    ]
    for label, values, color, maximum, box in panels:
        left, top, right, bottom = box
        draw.text((left, top - 36), label, fill=_rgb(COLORS["ink"]), font=_font(19, True))
        svg_body.append(_svg_text(left, top - 12, label, 19, bold=True))
        for tick in range(5):
            value = maximum * tick / 4
            y = bottom - (bottom - top) * tick / 4
            draw.line((left, y, right, y), fill=_rgb(COLORS["grid"]), width=2)
            draw.text((95, y - 10), f"{value:.0f}", fill=_rgb(COLORS["muted"]), font=_font(15))
            svg_body.extend(
                [
                    f'<line x1="{left}" y1="{y:.2f}" x2="{right}" y2="{y:.2f}" stroke="#{COLORS["grid"]}" stroke-width="2"/>',
                    _svg_text(95, y + 5, f"{value:.0f}", 15, color="muted"),
                ]
            )
        points = [
            (
                left + (right - left) * minute / max(minutes[-1], 1e-9),
                bottom - (bottom - top) * min(value, maximum) / maximum,
            )
            for minute, value in zip(minutes, values, strict=True)
        ]
        draw.line(points, fill=_rgb(COLORS[color]), width=4)
        svg_points = " ".join(f"{x:.2f},{y:.2f}" for x, y in points)
        svg_body.append(
            f'<polyline fill="none" stroke="#{COLORS[color]}" stroke-width="4" points="{svg_points}"/>'
        )
    draw.text((730, 777), "elapsed minutes", fill=_rgb(COLORS["muted"]), font=_font(17, True))
    svg_body.append(_svg_text(730, 799, "elapsed minutes", 17, bold=True, color="muted"))
    image.save(png, format="PNG", dpi=(180, 180), optimize=False)
    _write_svg(svg, svg_body, "Candidate B resource trace")
    return png, svg


def build_campaign_chart(output_dir: Path, facts: dict[str, Any]) -> tuple[Path, Path]:
    png = output_dir / "04_campaign_durability_proof.png"
    svg = output_dir / "04_campaign_durability_proof.svg"
    campaign = facts["campaign"]
    items = [
        ("1 remote launch", "single registered process identity", "purple"),
        (
            f"{campaign['remote_adoption_event_count']} adoptions",
            "successor worker reclaimed the same run",
            "sage",
        ),
        (
            f"{campaign['sealed_artifact_count']} sealed artifacts",
            "model, manifests, metrics, logs",
            "amber",
        ),
        (
            f"{campaign['budget']['actual']:.2f} GPU-hours",
            "actual charged; reservation returned to zero",
            "rose",
        ),
    ]
    image, draw = _canvas(
        "Durable Ponyo campaign smoke proof",
        "One training step proves orchestration semantics only; it is not a retrieval-quality result.",
    )
    svg_body = [
        _svg_text(88, 92, "Durable Ponyo campaign smoke proof", 42, bold=True),
        _svg_text(
            88,
            140,
            "One training step proves orchestration semantics only; it is not a retrieval-quality result.",
            23,
            color="muted",
        ),
        f'<rect x="84" y="172" width="1432" height="646" rx="20" fill="#{COLORS["soft"]}"/>',
    ]
    for index, (headline, detail, color) in enumerate(items):
        x = 120 + (index % 2) * 720
        y = 225 + (index // 2) * 270
        fill_key = f"{color}_light" if f"{color}_light" in COLORS else "paper"
        draw.rounded_rectangle(
            (x, y, x + 650, y + 210),
            radius=18,
            fill=_rgb(COLORS[fill_key]),
            outline=_rgb(COLORS[color]),
            width=3,
        )
        draw.text((x + 35, y + 35), headline, fill=_rgb(COLORS[color]), font=_font(31, True))
        draw.text((x + 35, y + 95), detail, fill=_rgb(COLORS["ink"]), font=_font(21))
        svg_body.extend(
            [
                f'<rect x="{x}" y="{y}" width="650" height="210" rx="18" fill="#{COLORS.get(color + "_light", COLORS["paper"])}" stroke="#{COLORS[color]}" stroke-width="3"/>',
                _svg_text(x + 35, y + 68, headline, 31, bold=True, color=color),
                _svg_text(x + 35, y + 125, detail, 21),
            ]
        )
    draw.text(
        (120, 750),
        "Promotion: not performed  |  Protected test: unopened  |  Campaign remains active",
        fill=_rgb(COLORS["muted"]),
        font=_font(20, True),
    )
    svg_body.append(
        _svg_text(
            120,
            775,
            "Promotion: not performed  |  Protected test: unopened  |  Campaign remains active",
            20,
            bold=True,
            color="muted",
        )
    )
    image.save(png, format="PNG", dpi=(180, 180), optimize=False)
    _write_svg(svg, svg_body, "Durable Ponyo campaign smoke proof")
    return png, svg


def build_charts(output_dir: Path, facts: dict[str, Any]) -> dict[str, dict[str, Path]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    return {
        "loss": dict(zip(("png", "svg"), build_loss_chart(output_dir, facts), strict=True)),
        "dev": dict(zip(("png", "svg"), build_dev_chart(output_dir, facts), strict=True)),
        "resources": dict(
            zip(("png", "svg"), build_resource_chart(output_dir, facts), strict=True)
        ),
        "campaign": dict(zip(("png", "svg"), build_campaign_chart(output_dir, facts), strict=True)),
    }


def write_tables(output_dir: Path, facts: dict[str, Any]) -> dict[str, Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    training = facts["training"]
    training_path = output_dir / "candidate_b_training_loss.csv"
    write_csv(
        training_path,
        ["step", "epoch", "loss", "rolling_mean_10", "learning_rate", "grad_norm"],
        zip(
            training["steps"],
            training["epoch_progress"],
            training["losses"],
            training["rolling_loss"],
            training["learning_rates"],
            training["grad_norms"],
            strict=True,
        ),
    )
    dev_path = output_dir / "strict_dev_comparison.csv"
    metric_keys = [
        "exact_chunk_mrr",
        "exact_chunk_recall_at_10",
        "local_window_mrr",
        "local_window_recall_at_10",
        "same_video_mrr",
        "same_video_recall_at_10",
        "median_query_latency_ms",
        "model_footprint_bytes",
    ]
    write_csv(
        dev_path,
        ["model", *metric_keys],
        (
            [name, *(metrics[key] for key in metric_keys)]
            for name, metrics in facts["models"].items()
        ),
    )
    campaign_path = output_dir / "campaign_smoke_proof.csv"
    campaign = facts["campaign"]
    campaign_rows = [
        ("attempt_status", campaign["attempt_status"]),
        ("single_remote_launch_event", campaign["single_remote_launch_event"]),
        ("remote_adoption_event_count", campaign["remote_adoption_event_count"]),
        ("claim_generation", campaign["claim_generation"]),
        ("successor_lease_owner", campaign["successor_lease_owner"]),
        ("sealed_artifact_count", campaign["sealed_artifact_count"]),
        ("actual_gpu_hours", campaign["budget"]["actual"]),
        ("reserved_gpu_hours", campaign["budget"]["reserved"]),
        ("promotion_performed", campaign["promotion_performed"]),
        ("protected_test_opened", campaign["protected_test_opened"]),
        ("smoke_loss_step", campaign["loss_points"][0]["step"]),
        ("smoke_loss", campaign["loss_points"][0]["value"]),
    ]
    write_csv(campaign_path, ["fact", "value"], campaign_rows)
    milestone_path = output_dir / "milestone_status.csv"
    write_csv(
        milestone_path,
        ["milestone", "status", "evidence"],
        ([row["milestone"], row["status"], row["evidence"]] for row in facts["milestones"]),
    )
    source_path = output_dir / "source_evidence.csv"
    write_csv(
        source_path,
        ["role", "name", "path", "size_bytes", "sha256"],
        (
            [row["role"], row["name"], row["path"], row["size_bytes"], row["sha256"]]
            for row in facts["source_hashes"]
        ),
    )
    return {
        "training": training_path,
        "dev": dev_path,
        "campaign": campaign_path,
        "milestones": milestone_path,
        "sources": source_path,
    }


def _fmt(value: float) -> str:
    return f"{value:.6f}"


def _decision_text(facts: dict[str, Any]) -> str:
    base = facts["models"]["Base Qwen 0.6B"]
    candidate = facts["models"]["Candidate B"]
    return (
        "Keep Base Qwen3-Embedding-0.6B as the release champion. Candidate B "
        f"completed all {facts['training']['optimizer_steps']} optimizer steps and improved "
        f"same-video MRR ({base['same_video_mrr']:.3f} to {candidate['same_video_mrr']:.3f}), "
        f"but exact MRR fell ({base['exact_chunk_mrr']:.3f} to {candidate['exact_chunk_mrr']:.3f}) "
        f"and local-window MRR fell ({base['local_window_mrr']:.3f} to {candidate['local_window_mrr']:.3f}). "
        "The 18-query dev set is characterization evidence, not a production-grade promotion gate."
    )


def write_markdown(path: Path, facts: dict[str, Any], charts: dict[str, dict[str, Path]]) -> None:
    relative = {
        name: values["png"].relative_to(path.parent).as_posix() for name, values in charts.items()
    }
    training = facts["training"]
    resources = facts["resources"]
    base = facts["models"]["Base Qwen 0.6B"]
    candidate_a = facts["models"]["Candidate A"]
    candidate_b = facts["models"]["Candidate B"]
    campaign = facts["campaign"]
    product = facts["product"]["metrics"]
    footprint_ratio = candidate_b["model_footprint_bytes"] / base["model_footprint_bytes"]
    milestone_lines = "\n".join(
        f"| {row['milestone']} | {row['status']} | {row['evidence']} |"
        for row in facts["milestones"]
    )
    text = f"""# {REPORT_TITLE}

{REPORT_SUBTITLE}.

Evidence cutoff: `{facts["evidence_cutoff"]}`.

## Decision

**Do not promote Candidate B.** {_decision_text(facts)}

The campaign smoke is a separate one-step run. It proves durable server-owned execution, restart adoption, artifact sealing, and budget accounting; it does **not** contribute to the model-quality decision.

## Full-run Candidate B findings

- Status: `{facts["training_manifest"]["status"]}`; {training["optimizer_steps"]} optimizer steps over {facts["training_manifest"]["train_pairs"]} rows and {training["epochs"]} epochs.
- Objective: `{training["loss"]}` with requested batch {training["batch_size_requested"]}, collision-safe realized mean {training["batch_size_realized_mean"]:.2f}, mini-batch {training["mini_batch_size"]}, and `{training["precision"]}` precision.
- Runtime: {training["result_metrics"]["train_runtime"]:.1f} seconds ({training["result_metrics"]["train_runtime"] / 60:.1f} minutes).
- Logged loss: first {training["first_loss"]:.4f}, last {training["last_loss"]:.4f}, minimum {training["minimum_loss"]:.4f}; trainer result loss {training["result_metrics"]["train_loss"]:.6f}.
- False-negative protection: `{training["positive_collision_group"]}` collision groups are excluded within each batch.

![Candidate B full-run loss]({relative["loss"]})

## Strict development comparison

All three models use the same physical `heldout-dev.jsonl`, restricted to 18 development queries. The frozen test is not part of this package.

| Model | Exact MRR | Exact R@10 | Local MRR | Local R@10 | Same-video MRR | Median query ms | Footprint bytes |
|---|---:|---:|---:|---:|---:|---:|---:|
| Base Qwen 0.6B | {_fmt(base["exact_chunk_mrr"])} | {_fmt(base["exact_chunk_recall_at_10"])} | {_fmt(base["local_window_mrr"])} | {_fmt(base["local_window_recall_at_10"])} | {_fmt(base["same_video_mrr"])} | {base["median_query_latency_ms"]:.3f} | {base["model_footprint_bytes"]:,} |
| Candidate A | {_fmt(candidate_a["exact_chunk_mrr"])} | {_fmt(candidate_a["exact_chunk_recall_at_10"])} | {_fmt(candidate_a["local_window_mrr"])} | {_fmt(candidate_a["local_window_recall_at_10"])} | {_fmt(candidate_a["same_video_mrr"])} | {candidate_a["median_query_latency_ms"]:.3f} | {candidate_a["model_footprint_bytes"]:,} |
| Candidate B | {_fmt(candidate_b["exact_chunk_mrr"])} | {_fmt(candidate_b["exact_chunk_recall_at_10"])} | {_fmt(candidate_b["local_window_mrr"])} | {_fmt(candidate_b["local_window_recall_at_10"])} | {_fmt(candidate_b["same_video_mrr"])} | {candidate_b["median_query_latency_ms"]:.3f} | {candidate_b["model_footprint_bytes"]:,} |

![Strict dev comparison]({relative["dev"]})

Candidate B improves broad recall and same-video ranking, but misses the exact and local-window MRR directions that matter for release selection. Its copied model footprint is {footprint_ratio:.2f}x the base footprint, so release serialization must also be corrected and revalidated before promotion is possible.

## Product compatibility fixture

Candidate B passed the product fixture: MRR {product["mrr"]:.3f}, Recall@1 {product["recallAt1"]:.3f}, nDCG@5 {product["ndcgAt5"]:.6f}, wrong-top rate {product["wrongTopRate"]:.3f}, and hard-negative-above-positive rate {product["hardNegativeAbovePositiveRate"]:.3f}. Because the fixture is saturated, this is compatibility evidence only and cannot select a champion.

## Ponyo resource envelope

Across {resources["sample_count"]} samples, available unified memory reached {resources["available_memory_floor_gib"]:.2f} GiB, swap grew {resources["swap_growth_gib"]:.2f} GiB, mean GPU utilization was {resources["gpu_utilization_mean_pct"]:.1f}%, and peak temperature was {resources["temperature_peak_c"]:.0f} C.

![Candidate B resource trace]({relative["resources"]})

## Durable campaign smoke proof

Attempt `{campaign["attempt_id"]}` completed after exactly one registered remote launch. A successor worker adopted the same remote identity {campaign["remote_adoption_event_count"]} times; the final claim generation is {campaign["claim_generation"]}. The campaign sealed {campaign["sealed_artifact_count"]} artifacts, charged {campaign["budget"]["actual"]:.2f} GPU-hours, returned reserved budget to {campaign["budget"]["reserved"]:.2f}, did not promote a model, and did not open the protected test.

The smoke has one loss point (step 1, loss {campaign["loss_points"][0]["value"]:.4f}). It is orchestration evidence only.

![Campaign durability proof]({relative["campaign"]})

The campaign export manifest contains {len(facts["campaign_export"]["verified_files"])} files; every declared file was rehashed and matched its recorded SHA-256 and size. Its own `quality_findings_available` flag is `{str(facts["campaign_export"]["quality_findings_available"]).lower()}`, which is correct for this smoke.

## Milestone audit

| Milestone | Status | Evidence |
|---|---|---|
{milestone_lines}

## Flags and limitations

1. **No promotion:** Base Qwen remains champion because Candidate B regressed exact and local-window MRR.
2. **Small dev set:** 18 queries across three videos are not enough for a stable production promotion gate; expand and adjudicate the set before causal conclusions.
3. **Footprint failure:** Candidate B is {footprint_ratio:.2f}x the base model footprint in the copied evaluation artifacts. Produce a BF16 release serialization and prove embedding/ranking equivalence.
4. **Saturated product fixture:** perfect top-rank compatibility does not distinguish candidates.
5. **Smoke/quality boundary:** the one-step campaign run proves orchestration only.
6. **System ablations missing:** no BM25, dense/BM25 RRF, or independent reranker scores are present.
7. **Human labeling missing:** no 30-query double-label pilot or inter-annotator agreement artifact is present.
8. **Operational closure missing:** this package does not prove a persistent resident worker service or a single-campaign API/CLI/MCP/live-canvas trace.

## Recommended next steps

1. Expand the video-disjoint development set first; double-label at least the planned 30-query pilot and record disagreement/adjudication.
2. Evaluate dense base/Candidate B, BM25, dense+BM25 RRF, and a separately pinned reranker on the identical query/corpus lineage. Report per-query ranks, nDCG, MRR, Recall, hard-negative wins, and paired confidence intervals.
3. Serialize Candidate B in BF16, measure directory/model size, and verify embedding cosine/ranking equivalence before any serving benchmark.
4. Run lower-pressure Candidate B ablations (one epoch/checkpoint selection and negative-quality changes) before increasing batch or duration again.
5. Install and exercise the resident campaign worker through restart/reboot, then capture one campaign end to end through REST, CLI/MCP, and the live canvas.
6. Keep the frozen test closed until a candidate passes the expanded development gate; promotion must remain explicit and audited.

## Reproducibility package

The `tables/` directory contains the full 84-point loss series, strict dev metrics, smoke facts, milestone matrix, and source hashes. The `charts/` directory contains PNG and SVG versions of every figure. `report_manifest.json` records source/output hashes after generation.
"""
    path.write_text(text, encoding="utf-8")


def _set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    italic: bool | None = None,
) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def _set_table_geometry(table, widths: Sequence[int]) -> None:
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag, value in (("tblW", sum(widths)), ("tblInd", 120)):
        node = tbl_pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            margins = tc_pr.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for margin, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = margins.find(qn(f"w:{margin}"))
                if node is None:
                    node = OxmlElement(f"w:{margin}")
                    margins.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if table.rows:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            node = OxmlElement("w:tblHeader")
            node.set(qn("w:val"), "true")
            tr_pr.append(node)


def _configure_docx(doc) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = (
        Inches(1)
    )
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(COLORS["ink"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for level, (size, before, after, color) in {
        1: (16, 16, 8, COLORS["purple"]),
        2: (13, 12, 6, COLORS["purple"]),
        3: (12, 8, 4, COLORS["ink"]),
    }.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run_font(
        header.add_run("MEMEXAI  |  EMBEDDING AUTOMATION IMPLEMENTATION REPORT"),
        size=8.5,
        bold=True,
        color=COLORS["muted"],
    )
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_run_font(footer.add_run("Page "), size=9, color=COLORS["muted"])
    for text in ("PAGE", "NUMPAGES"):
        if text == "NUMPAGES":
            _set_run_font(footer.add_run(" of "), size=9, color=COLORS["muted"])
        run = footer.add_run()
        _set_run_font(run, size=9, color=COLORS["muted"])
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = text
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        display = OxmlElement("w:t")
        display.text = "1"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for element in (begin, instruction, separate, display, end):
            run._r.append(element)


def _docx_body(doc, text: str, *, bold_lead: str | None = None) -> None:
    from docx.shared import Pt

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_lead and text.startswith(bold_lead):
        _set_run_font(paragraph.add_run(bold_lead), bold=True, color=COLORS["ink"])
        _set_run_font(paragraph.add_run(text[len(bold_lead) :]), color=COLORS["ink"])
    else:
        _set_run_font(paragraph.add_run(text), color=COLORS["ink"])


def _docx_heading(doc, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    _set_run_font(
        paragraph.add_run(text), bold=True, color=COLORS["purple"] if level < 3 else COLORS["ink"]
    )


def _docx_table(
    doc, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int]
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, COLORS["soft"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        _set_run_font(p.add_run(header), size=8.5, bold=True, color=COLORS["ink"])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            p = cells[index].paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if index in (0, len(row) - 1) else WD_ALIGN_PARAGRAPH.CENTER
            )
            p.paragraph_format.space_after = Pt(0)
            _set_run_font(p.add_run(str(value)), size=8.5, color=COLORS["ink"])
    _set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def _docx_callout(doc, label: str, text: str, *, fill: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    _set_run_font(paragraph.add_run(f"{label}: "), bold=True, color=COLORS["purple"])
    _set_run_font(paragraph.add_run(text), color=COLORS["ink"])
    _set_table_geometry(table, [9360])
    doc.add_paragraph().paragraph_format.space_after = 1


def _docx_figure(doc, path: Path, caption: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    inline = paragraph.add_run().add_picture(str(path), width=Inches(6.25))
    inline._inline.docPr.set("descr", caption)
    inline._inline.docPr.set("title", caption)
    caption_p = doc.add_paragraph(style="Caption")
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_before = Pt(2)
    caption_p.paragraph_format.space_after = Pt(8)
    _set_run_font(caption_p.add_run(caption), size=9, italic=True, color=COLORS["muted"])


def normalize_docx_zip(path: Path) -> None:
    """Normalize ZIP order/timestamps so identical inputs produce identical DOCX bytes."""
    temp = path.with_suffix(".normalized.docx")
    with (
        zipfile.ZipFile(path, "r") as source,
        zipfile.ZipFile(temp, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as target,
    ):
        for name in sorted(source.namelist()):
            info = zipfile.ZipInfo(name, (1980, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            target.writestr(info, source.read(name))
    temp.replace(path)


def build_docx(path: Path, facts: dict[str, Any], charts: dict[str, dict[str, Path]]) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    _configure_docx(doc)
    # memo_masthead header pattern, resolved through the standard_business_brief preset
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    _set_run_font(
        kicker.add_run("TECHNICAL IMPLEMENTATION REPORT"),
        size=10.5,
        bold=True,
        color=COLORS["sage"],
    )
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    _set_run_font(title.add_run(REPORT_TITLE), size=25, bold=True, color=COLORS["ink"])
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    _set_run_font(subtitle.add_run(REPORT_SUBTITLE), size=12.5, color=COLORS["muted"])
    metadata = [
        ["Evidence cutoff", facts["evidence_cutoff"]],
        [
            "Full training run",
            f"Candidate B; {facts['training']['optimizer_steps']} steps; {facts['training']['epochs']} epochs",
        ],
        ["Campaign proof", f"{facts['campaign']['campaign_id']}; one-step orchestration smoke"],
        ["Release decision", "Do not promote Candidate B; Base Qwen remains champion"],
    ]
    _docx_table(doc, ["Field", "Value"], metadata, [2100, 7260])
    _docx_callout(doc, "Decision", _decision_text(facts), fill=COLORS["amber_light"])
    doc.add_page_break()

    training = facts["training"]
    _docx_heading(doc, "Full-run Candidate B findings", 1)
    _docx_body(
        doc,
        f"Candidate B completed {training['optimizer_steps']} cached-MNRL optimizer steps over {facts['training_manifest']['train_pairs']} rows and {training['epochs']} epochs in {training['result_metrics']['train_runtime'] / 60:.1f} minutes. Requested batch {training['batch_size_requested']} was collision- and memory-limited to a realized mean {training['batch_size_realized_mean']:.2f}, with mini-batch {training['mini_batch_size']} and {training['precision']} precision.",
    )
    _docx_figure(
        doc,
        charts["loss"]["png"],
        "Figure 1. Candidate B full-run loss; the one-step campaign smoke is excluded.",
    )
    _docx_body(
        doc,
        f"Logged loss moved from {training['first_loss']:.4f} to {training['last_loss']:.4f}; minimum logged loss was {training['minimum_loss']:.4f}. Loss reduction proves the objective was optimized, not that retrieval improved.",
    )

    doc.add_page_break()
    _docx_heading(doc, "Strict development comparison", 1)
    _docx_body(
        doc,
        "Base, Candidate A, and Candidate B were scored against the same 18-row dev-only heldout-dev.jsonl and 2,018-chunk corpus. The frozen test is not part of this evidence package.",
    )
    _docx_figure(
        doc,
        charts["dev"]["png"],
        "Figure 2. Base and trained candidates on the strict held-out development split.",
    )
    metric_rows = []
    for name, metrics in facts["models"].items():
        metric_rows.append(
            [
                name,
                f"{metrics['exact_chunk_mrr']:.3f}",
                f"{metrics['exact_chunk_recall_at_10']:.3f}",
                f"{metrics['local_window_mrr']:.3f}",
                f"{metrics['same_video_mrr']:.3f}",
            ]
        )
    _docx_table(
        doc,
        ["Model", "Exact MRR", "Exact R@10", "Local MRR", "Video MRR"],
        metric_rows,
        [2800, 1500, 1600, 1700, 1760],
    )
    base = facts["models"]["Base Qwen 0.6B"]
    candidate = facts["models"]["Candidate B"]
    ratio = candidate["model_footprint_bytes"] / base["model_footprint_bytes"]
    _docx_callout(
        doc,
        "Promotion gate",
        f"Candidate B improves same-video MRR but regresses exact and local-window MRR. Its copied footprint is {ratio:.2f}x base. Keep base as champion and treat this 18-query result as characterization, not a universal claim.",
        fill=COLORS["amber_light"],
    )
    product = facts["product"]["metrics"]
    _docx_body(
        doc,
        f"The product fixture passed (MRR {product['mrr']:.3f}, Recall@1 {product['recallAt1']:.3f}, nDCG@5 {product['ndcgAt5']:.6f}), but is saturated and supplies compatibility rather than model-selection evidence.",
    )

    doc.add_page_break()
    _docx_heading(doc, "Ponyo resource envelope", 1)
    _docx_figure(
        doc,
        charts["resources"]["png"],
        "Figure 3. Unified-memory availability and GPU utilization across the full run.",
    )
    resources = facts["resources"]
    _docx_body(
        doc,
        f"Across {resources['sample_count']} samples, available unified memory reached {resources['available_memory_floor_gib']:.2f} GiB, swap grew {resources['swap_growth_gib']:.2f} GiB, mean GPU utilization was {resources['gpu_utilization_mean_pct']:.1f}%, and peak temperature was {resources['temperature_peak_c']:.0f} C.",
    )

    doc.add_page_break()
    _docx_heading(doc, "Durable campaign smoke proof", 1)
    _docx_callout(
        doc,
        "Evidence boundary",
        "This independent one-step run proves orchestration semantics only. It is not used in the Candidate B retrieval decision.",
        fill=COLORS["purple_light"],
    )
    _docx_figure(
        doc,
        charts["campaign"]["png"],
        "Figure 4. Single launch, successor adoption, sealing, and budget proof.",
    )
    campaign = facts["campaign"]
    _docx_body(
        doc,
        f"Attempt {campaign['attempt_id']} completed after one registered launch. A successor adopted the exact remote identity {campaign['remote_adoption_event_count']} times. The campaign sealed {campaign['sealed_artifact_count']} artifacts, charged {campaign['budget']['actual']:.2f} GPU-hours, returned reservation to zero, performed no promotion, and left the protected test unopened.",
    )
    _docx_body(
        doc,
        f"The campaign export contains {len(facts['campaign_export']['verified_files'])} files whose sizes and SHA-256 values were reverified. The export correctly reports quality_findings_available=false for this smoke.",
    )

    doc.add_page_break()
    _docx_heading(doc, "Milestone audit", 1)
    milestone_doc_rows = [
        [row["milestone"], row["status"], row["evidence"]] for row in facts["milestones"]
    ]
    _docx_table(doc, ["Milestone", "Status", "Evidence"], milestone_doc_rows, [3300, 1400, 4660])

    doc.add_page_break()
    _docx_heading(doc, "Flags and limitations", 1)
    flags = [
        "Base remains champion; Candidate B regressed exact and local-window MRR.",
        "The 18-query, three-video dev set is too small for a production promotion gate.",
        f"Candidate B's copied model footprint is {ratio:.2f}x base; BF16 serialization and equivalence proof are required.",
        "The product fixture is saturated and cannot select among candidates.",
        "The one-step campaign run proves orchestration, not retrieval quality.",
        "No scored BM25/RRF/reranker artifact or 30-query labeling pilot is present.",
        "No persistent resident worker or integrated API/CLI/MCP/live-canvas trace is present.",
    ]
    for item in flags:
        p = doc.add_paragraph(style="List Bullet")
        _set_run_font(p.add_run(item), color=COLORS["ink"])
    _docx_heading(doc, "Recommended next steps", 1)
    next_steps = [
        "Expand and double-label the video-disjoint development set before another promotion decision.",
        "Run dense base/Candidate B, BM25, dense+BM25 RRF, and a pinned independent reranker on identical lineage with per-query outputs and paired intervals.",
        "Serialize Candidate B in BF16 and prove embedding/ranking equivalence plus serving footprint.",
        "Test lower-pressure one-epoch/checkpoint and negative-quality ablations before increasing batch or duration.",
        "Install and reboot-test the resident worker, then capture one campaign through REST, CLI/MCP, and the live canvas.",
        "Keep the frozen test closed until a future candidate passes the expanded development gate.",
    ]
    for item in next_steps:
        p = doc.add_paragraph(style="List Number")
        _set_run_font(p.add_run(item), color=COLORS["ink"])

    _docx_heading(doc, "Reproducibility contract", 1)
    _docx_body(
        doc,
        "The companion CSV tables contain the complete loss series, strict development metrics, smoke proof, milestone audit, and source hashes. Every figure is available as both PNG and SVG. report_manifest.json records the output hashes.",
    )
    source_rows = [[row["role"], row["sha256"]] for row in facts["source_hashes"]]
    _docx_table(doc, ["Evidence role", "SHA-256"], source_rows, [3100, 6260])

    fixed_dt = datetime.fromisoformat(facts["evidence_cutoff"].replace("Z", "+00:00"))
    if fixed_dt.tzinfo is not None:
        fixed_dt = fixed_dt.astimezone(timezone.utc).replace(tzinfo=None)
    doc.core_properties.title = REPORT_TITLE
    doc.core_properties.subject = REPORT_SUBTITLE
    doc.core_properties.author = "BashGym / MemexAI"
    doc.core_properties.keywords = "MemexAI, embeddings, automation, Ponyo, evaluation"
    doc.core_properties.created = fixed_dt
    doc.core_properties.modified = fixed_dt
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    normalize_docx_zip(path)


def build_pdf(path: Path, facts: dict[str, Any], charts: dict[str, dict[str, Path]]) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Image,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="ReportTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=24,
            leading=28,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ReportCaption",
            parent=styles["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=8.5,
            leading=11,
            alignment=1,
            textColor=colors.HexColor("#667085"),
            spaceAfter=10,
        )
    )
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        title=REPORT_TITLE,
        author="BashGym / MemexAI",
        invariant=1,
    )
    body = styles["ReportBody"]
    story: list[Any] = [
        Paragraph("TECHNICAL IMPLEMENTATION REPORT", styles["Heading3"]),
        Paragraph(REPORT_TITLE, styles["ReportTitle"]),
        Paragraph(REPORT_SUBTITLE, styles["Heading2"]),
        Spacer(1, 0.15 * inch),
        Paragraph(f"<b>Decision:</b> {_decision_text(facts)}", body),
        Paragraph(
            "The one-step campaign smoke is separate orchestration evidence and is never used as model-quality evidence.",
            body,
        ),
        PageBreak(),
        Paragraph("Full-run Candidate B findings", styles["Heading1"]),
        Paragraph(
            f"Candidate B completed {facts['training']['optimizer_steps']} optimizer steps over {facts['training_manifest']['train_pairs']} rows in {facts['training']['result_metrics']['train_runtime'] / 60:.1f} minutes. Logged loss moved from {facts['training']['first_loss']:.4f} to {facts['training']['last_loss']:.4f}.",
            body,
        ),
        Image(str(charts["loss"]["png"]), width=6.5 * inch, height=3.66 * inch),
        Paragraph(
            "Figure 1. Complete Candidate B loss series; smoke excluded.", styles["ReportCaption"]
        ),
        PageBreak(),
        Paragraph("Strict development comparison", styles["Heading1"]),
        Image(str(charts["dev"]["png"]), width=6.5 * inch, height=3.66 * inch),
        Paragraph(
            "Figure 2. Base and trained candidates on 18 dev-only queries.", styles["ReportCaption"]
        ),
    ]
    dev_data = [["Model", "Exact MRR", "Exact R@10", "Local MRR", "Video MRR"]]
    for name, metrics in facts["models"].items():
        dev_data.append(
            [
                name,
                f"{metrics['exact_chunk_mrr']:.3f}",
                f"{metrics['exact_chunk_recall_at_10']:.3f}",
                f"{metrics['local_window_mrr']:.3f}",
                f"{metrics['same_video_mrr']:.3f}",
            ]
        )
    table = Table(
        dev_data,
        colWidths=[2.0 * inch, 1.05 * inch, 1.05 * inch, 1.15 * inch, 1.15 * inch],
        repeatRows=1,
    )
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D9DEE7")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend(
        [
            table,
            Paragraph(
                "Candidate B improves broad recall/same-video ranking but misses exact and local-window MRR. The product fixture passes but is saturated.",
                body,
            ),
            PageBreak(),
            Paragraph("Ponyo resource envelope", styles["Heading1"]),
            Image(str(charts["resources"]["png"]), width=6.5 * inch, height=3.66 * inch),
            Paragraph("Figure 3. Full-run unified-memory and GPU trace.", styles["ReportCaption"]),
            PageBreak(),
            Paragraph("Durable campaign smoke proof", styles["Heading1"]),
            Paragraph(
                "This independent one-step run proves launch ownership, restart adoption, sealing, and accounting only.",
                body,
            ),
            Image(str(charts["campaign"]["png"]), width=6.5 * inch, height=3.66 * inch),
            Paragraph("Figure 4. Durable orchestration proof.", styles["ReportCaption"]),
        ]
    )
    campaign = facts["campaign"]
    story.extend(
        [
            Paragraph(
                f"One launch; {campaign['remote_adoption_event_count']} adoptions; {campaign['sealed_artifact_count']} sealed artifacts; {campaign['budget']['actual']:.2f} GPU-hours; no promotion; protected test unopened.",
                body,
            ),
            PageBreak(),
            Paragraph("Milestone audit", styles["Heading1"]),
        ]
    )
    milestone_data = [["Milestone", "Status", "Evidence"]] + [
        [Paragraph(row["milestone"], body), row["status"], Paragraph(row["evidence"], body)]
        for row in facts["milestones"]
    ]
    milestone_table = Table(
        milestone_data, colWidths=[2.2 * inch, 1.05 * inch, 3.25 * inch], repeatRows=1
    )
    milestone_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D9DEE7")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend(
        [
            milestone_table,
            PageBreak(),
            Paragraph("Flags and next steps", styles["Heading1"]),
            Paragraph(
                "<b>Flags:</b> small dev set; Candidate B footprint near 2x base; saturated fixture; no BM25/RRF/reranker benchmark; no double-label pilot; no resident-service or live cross-surface proof.",
                body,
            ),
            Paragraph(
                "<b>Next:</b> expand/adjudicate dev, run system ablations on identical lineage, serialize BF16 with equivalence proof, test lower-pressure training, install/reboot-test the resident worker, and capture one campaign across REST, CLI/MCP, and the live canvas.",
                body,
            ),
            Paragraph(
                "Keep the frozen test closed until a future candidate passes the expanded development gate.",
                body,
            ),
        ]
    )
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.build(story)


def write_report_manifest(output_dir: Path, facts: dict[str, Any]) -> Path:
    path = output_dir / "report_manifest.json"
    outputs = []
    for target in sorted(output_dir.rglob("*")):
        if target.is_file() and target != path and "qa" not in target.parts:
            outputs.append(
                {
                    "path": target.relative_to(output_dir).as_posix(),
                    "size_bytes": target.stat().st_size,
                    "sha256": sha256_file(target),
                }
            )
    payload = {
        "schema_version": "memexai_embedding_automation_report.v1",
        "report_title": REPORT_TITLE,
        "evidence_cutoff": facts["evidence_cutoff"],
        "source_evidence": facts["source_hashes"],
        "outputs": outputs,
        "quality_decision": "do_not_promote_candidate_b",
        "smoke_evidence_boundary": "orchestration_only",
    }
    path.write_text(json.dumps(payload, indent=2, sort_keys=True) + "\n", encoding="utf-8")
    return path


def build_report(inputs: ReportInputs, output_dir: Path) -> dict[str, Any]:
    facts = load_report_facts(inputs)
    charts = build_charts(output_dir / "charts", facts)
    tables = write_tables(output_dir / "tables", facts)
    markdown = output_dir / f"{REPORT_STEM}.md"
    docx = output_dir / f"{REPORT_STEM}.docx"
    pdf = output_dir / f"{REPORT_STEM}.pdf"
    write_markdown(markdown, facts, charts)
    build_docx(docx, facts, charts)
    build_pdf(pdf, facts, charts)
    manifest = write_report_manifest(output_dir, facts)
    return {
        "facts": facts,
        "charts": charts,
        "tables": tables,
        "markdown": markdown,
        "docx": docx,
        "pdf": pdf,
        "manifest": manifest,
    }


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--candidate-run-dir", type=Path, required=True)
    parser.add_argument("--base-dev-manifest", type=Path, required=True)
    parser.add_argument("--candidate-a-dev-manifest", type=Path, required=True)
    parser.add_argument("--campaign-evidence", type=Path, required=True)
    parser.add_argument("--campaign-export-manifest", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    result = build_report(
        ReportInputs(
            candidate_run_dir=args.candidate_run_dir,
            base_dev_manifest=args.base_dev_manifest,
            candidate_a_dev_manifest=args.candidate_a_dev_manifest,
            campaign_evidence=args.campaign_evidence,
            campaign_export_manifest=args.campaign_export_manifest,
        ),
        args.output_dir,
    )
    print(result["manifest"].resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
