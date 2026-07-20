#!/usr/bin/env python3
"""Extend the MemexAI experiment report with the positive-aware campaign findings."""

from __future__ import annotations

import argparse
import csv
import json
import math
import random
import statistics
from pathlib import Path
from typing import Any


def report_facts(manifest: dict[str, Any]) -> dict[str, Any]:
    stats = manifest["statistics"]
    rows = int(stats["rows"])
    ready = int(stats["ready_rows"])
    grouped = sum(
        int(count) for size, count in stats.get("positive_group_sizes", {}).items() if int(size) > 1
    )
    return {
        "rows": rows,
        "ready_rows": ready,
        "insufficient_rows": int(stats["insufficient_safe_negatives"]),
        "coverage": ready / rows if rows else 0.0,
        "selected_negatives": sum(map(int, stats.get("negative_sources", {}).values())),
        "excluded_candidates": sum(map(int, stats.get("exclusion_reasons", {}).values())),
        "grouped_positive_rows": grouped,
        "negative_sources": stats.get("negative_sources", {}),
        "positive_group_sizes": stats.get("positive_group_sizes", {}),
        "exclusion_reasons": stats.get("exclusion_reasons", {}),
    }


def load_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def load_jsonl(path: Path) -> list[dict[str, Any]]:
    return [
        json.loads(line) for line in path.read_text(encoding="utf-8").splitlines() if line.strip()
    ]


def rolling_mean(values: list[float], window: int = 15) -> list[float]:
    if window < 1:
        raise ValueError("rolling window must be >= 1")
    return [
        statistics.fmean(values[max(0, index - window + 1) : index + 1])
        for index in range(len(values))
    ]


def summarize_training_metrics(
    manifest: dict[str, Any], metrics_rows: list[dict[str, Any]]
) -> dict[str, Any]:
    loss_rows = sorted(
        (row for row in metrics_rows if "loss" in row and "step" in row),
        key=lambda row: int(row["step"]),
    )
    if not loss_rows:
        raise ValueError("training metrics contain no per-step loss rows")
    steps = [int(row["step"]) for row in loss_rows]
    if len(steps) != len(set(steps)):
        raise ValueError("training metrics contain duplicate loss steps")
    losses = [float(row["loss"]) for row in loss_rows]
    expected_steps = int(manifest["training"]["optimizer_steps"])
    if steps != list(range(1, expected_steps + 1)):
        raise ValueError(
            f"training loss steps are incomplete: expected 1..{expected_steps}, got {steps[:1]}..{steps[-1:]}"
        )
    batches_per_epoch = int(manifest["training"]["collision_safe_batches_per_epoch"])

    def describe(values: list[float]) -> dict[str, float | int]:
        ordered = sorted(values)
        p90_index = max(0, math.ceil(0.90 * len(ordered)) - 1)
        return {
            "count": len(values),
            "mean": statistics.fmean(values),
            "median": statistics.median(values),
            "p90": ordered[p90_index],
            "maximum": max(values),
            "above_1_count": sum(value > 1.0 for value in values),
        }

    return {
        "loss_steps": steps,
        "loss_values": losses,
        "loss_rolling_mean_15": rolling_mean(losses, 15),
        "loss_overall": describe(losses),
        "loss_epoch_1": describe(losses[:batches_per_epoch]),
        "loss_epoch_2": describe(losses[batches_per_epoch:]),
        "epoch_boundary_step": batches_per_epoch,
        "learning_rates": [float(row["learning_rate"]) for row in loss_rows],
        "training_result": manifest["training"]["result_metrics"],
        "training_manifest": manifest,
    }


def load_resource_samples(path: Path) -> list[dict[str, float]]:
    with path.open("r", encoding="utf-8", newline="") as handle:
        rows = [{key: float(value) for key, value in row.items()} for row in csv.DictReader(handle)]
    if not rows:
        raise ValueError("resource sample CSV is empty")
    return rows


def summarize_resources(rows: list[dict[str, float]]) -> dict[str, Any]:
    available = [row["mem_available_bytes"] / (1024**3) for row in rows]
    swap = [row["swap_used_bytes"] / (1024**3) for row in rows]
    gpu = [row["gpu_util_pct"] for row in rows]
    temperatures = [row["temp_c"] for row in rows]
    start = rows[0]["epoch_utc"]
    return {
        "resource_minutes": [(row["epoch_utc"] - start) / 60 for row in rows],
        "available_memory_gib": available,
        "gpu_utilization_pct": gpu,
        "available_memory_floor_gib": min(available),
        "swap_start_gib": swap[0],
        "swap_peak_gib": max(swap),
        "swap_growth_gib": max(swap) - swap[0],
        "gpu_utilization_mean_pct": statistics.fmean(gpu),
        "temperature_peak_c": max(temperatures),
        "resource_sample_count": len(rows),
    }


def paired_bootstrap_mrr_deltas(
    base_rows: list[dict[str, Any]],
    candidate_rows: list[dict[str, Any]],
    *,
    samples: int = 10_000,
    seed: int = 20260712,
) -> dict[str, dict[str, float]]:
    base = {str(row["eval_id"]): row for row in base_rows}
    candidate = {str(row["eval_id"]): row for row in candidate_rows}
    if set(base) != set(candidate):
        raise ValueError("base and candidate dev rows do not contain identical eval IDs")
    ids = sorted(base)
    if not ids:
        raise ValueError("dev comparison rows are empty")
    result: dict[str, dict[str, float]] = {}
    for offset, (label, rank_key) in enumerate(
        (
            ("exact_mrr", "positive_rank_exact"),
            ("local_window_mrr", "positive_rank_local_window"),
            ("same_video_mrr", "positive_rank_same_video"),
        )
    ):
        per_query = [
            (1.0 / float(candidate[eval_id][rank_key])) - (1.0 / float(base[eval_id][rank_key]))
            for eval_id in ids
        ]
        rng = random.Random(seed + offset)
        bootstrapped = sorted(
            statistics.fmean(per_query[rng.randrange(len(per_query))] for _ in per_query)
            for _ in range(samples)
        )
        result[label] = {
            "delta": statistics.fmean(per_query),
            "ci95_low": bootstrapped[math.floor(0.025 * (samples - 1))],
            "ci95_high": bootstrapped[math.floor(0.975 * (samples - 1))],
            "rows": len(ids),
            "bootstrap_samples": samples,
        }
    return result


def add_run_facts(
    facts: dict[str, Any],
    *,
    training_manifest_path: Path,
    training_metrics_path: Path,
    resource_samples_path: Path,
    base_dev_manifest_path: Path,
    candidate_dev_manifest_path: Path,
    base_dev_rows_path: Path,
    candidate_dev_rows_path: Path,
    product_eval_path: Path,
) -> dict[str, Any]:
    manifest = load_json(training_manifest_path)
    facts.update(summarize_training_metrics(manifest, load_jsonl(training_metrics_path)))
    facts.update(summarize_resources(load_resource_samples(resource_samples_path)))
    base_dev = load_json(base_dev_manifest_path)["runs"]["memexai_youtube"]["metrics"]["overall"]
    candidate_dev = load_json(candidate_dev_manifest_path)["runs"]["memexai_youtube"]["metrics"][
        "overall"
    ]
    facts["base_dev"] = base_dev
    facts["candidate_dev"] = candidate_dev
    facts["dev_bootstrap"] = paired_bootstrap_mrr_deltas(
        load_jsonl(base_dev_rows_path), load_jsonl(candidate_dev_rows_path)
    )
    facts["product_eval"] = load_json(product_eval_path)
    facts["promotion_status"] = "do_not_promote_dev_regression"
    facts["frozen_test_status"] = "unopened_after_dev_gate_failure"
    return facts


def build_charts(output_dir: Path, facts: dict[str, Any]) -> dict[str, Path]:
    from PIL import ImageDraw

    from scripts.memexai.build_embedding_experiment_report import (
        COLORS,
        chart_canvas,
        font,
        rgb,
    )

    output_dir.mkdir(parents=True, exist_ok=True)
    coverage_path = output_dir / "07_positive_aware_dataset_coverage.png"
    image, draw = chart_canvas(
        "Positive-aware training coverage",
        "Real702 queries with at least three safe explicit negatives after false-negative filtering.",
    )
    left, right, top, bottom = 180, 1420, 330, 500
    draw.rounded_rectangle((left, top, right, bottom), radius=22, fill=rgb(COLORS["grid"]))
    ready_width = int((right - left) * facts["coverage"])
    draw.rounded_rectangle(
        (left, top, left + ready_width, bottom), radius=22, fill=rgb(COLORS["sage"])
    )
    draw.text(
        (left, 250),
        f"{facts['ready_rows']} ready / {facts['rows']} total ({facts['coverage']:.1%})",
        fill=rgb(COLORS["ink"]),
        font=font(34, True),
    )
    draw.text(
        (left, 555),
        f"{facts['insufficient_rows']} rows remain pair-only diagnostics; they are not silently trained.",
        fill=rgb(COLORS["muted"]),
        font=font(25),
    )
    image.save(coverage_path, dpi=(180, 180))

    source_path = output_dir / "08_negative_sources_and_exclusions.png"
    image, draw = chart_canvas(
        "Where the hard negatives came from",
        "Selected negatives are wrong-video candidates; unreviewed same-video candidates stay excluded.",
    )
    panels = [
        ("Selected", facts["negative_sources"], COLORS["sage"]),
        ("Excluded", facts["exclusion_reasons"], COLORS["amber"]),
    ]
    for panel_index, (title, values, color) in enumerate(panels):
        x0 = 110 + panel_index * 740
        draw.text((x0, 230), title, fill=rgb(COLORS["ink"]), font=font(31, True))
        total = max(sum(map(int, values.values())), 1)
        for row_index, (label, value) in enumerate(
            sorted(values.items(), key=lambda item: -int(item[1]))
        ):
            if row_index >= 4:
                break
            y = 315 + row_index * 112
            width = int(520 * int(value) / total)
            draw.text(
                (x0, y), label.replace("_", " "), fill=rgb(COLORS["ink"]), font=font(21, True)
            )
            draw.rounded_rectangle(
                (x0, y + 35, x0 + 520, y + 75), radius=8, fill=rgb(COLORS["grid"])
            )
            draw.rounded_rectangle((x0, y + 35, x0 + width, y + 75), radius=8, fill=rgb(color))
            draw.text((x0 + 535, y + 41), str(value), fill=rgb(COLORS["ink"]), font=font(19, True))
    image.save(source_path, dpi=(180, 180))

    memory_path = output_dir / "09_expanded_batch_memory_guard.png"
    image, draw = chart_canvas(
        "Why the first Ponyo batch probe was rejected",
        "Explicit negatives expand the physical query batch before the encoder forward pass.",
    )
    bars = [
        ("Batch 16 probe", 80, COLORS["amber"]),
        ("New hard ceiling", 64, COLORS["purple"]),
        ("Batch 4 run", 20, COLORS["sage"]),
    ]
    maximum = 90
    for row_index, (label, value, color) in enumerate(bars):
        y = 280 + row_index * 150
        draw.text((150, y), label, fill=rgb(COLORS["ink"]), font=font(25, True))
        draw.rounded_rectangle((500, y, 1370, y + 58), radius=10, fill=rgb(COLORS["grid"]))
        width = int(870 * value / maximum)
        draw.rounded_rectangle((500, y, 500 + width, y + 58), radius=10, fill=rgb(color))
        draw.text((1390, y + 12), str(value), fill=rgb(COLORS["ink"]), font=font(23, True))
    draw.text(
        (150, 745),
        "Actual lane count = batch × (query + positive + 3 fixed explicit negatives).",
        fill=rgb(COLORS["muted"]),
        font=font(22),
    )
    image.save(memory_path, dpi=(180, 180))

    def draw_line_panel(
        draw: ImageDraw.ImageDraw,
        values: list[float],
        *,
        box: tuple[int, int, int, int],
        color: str,
        maximum: float,
        label: str,
        x_values: list[float] | None = None,
        epoch_boundary: int | None = None,
        show_x_ticks: bool = False,
    ) -> None:
        left, top, right, bottom = box
        draw.text((left, top - 34), label, fill=rgb(COLORS["ink"]), font=font(21, True))
        for tick in range(5):
            value = maximum * tick / 4
            y = bottom - int((bottom - top) * tick / 4)
            draw.line((left, y, right, y), fill=rgb(COLORS["grid"]), width=2)
            draw.text((left - 72, y - 11), f"{value:.2f}", fill=rgb(COLORS["muted"]), font=font(16))
        xs = x_values or [float(index) for index in range(len(values))]
        minimum_x, maximum_x = min(xs), max(xs)
        points = []
        for x_value, value in zip(xs, values, strict=True):
            x = left + int((x_value - minimum_x) / max(maximum_x - minimum_x, 1) * (right - left))
            y = bottom - int(min(value, maximum) / max(maximum, 1e-12) * (bottom - top))
            points.append((x, y))
        if len(points) > 1:
            draw.line(points, fill=rgb(color), width=3, joint="curve")
        if show_x_ticks:
            for tick in range(5):
                x = left + int((right - left) * tick / 4)
                value = minimum_x + (maximum_x - minimum_x) * tick / 4
                label_text = f"{value:.0f}"
                bbox = draw.textbbox((0, 0), label_text, font=font(15))
                draw.text(
                    (x - (bbox[2] - bbox[0]) / 2, bottom + 4),
                    label_text,
                    fill=rgb(COLORS["muted"]),
                    font=font(15),
                )
        if epoch_boundary is not None:
            boundary_x = left + int(
                (epoch_boundary - minimum_x) / max(maximum_x - minimum_x, 1) * (right - left)
            )
            draw.line((boundary_x, top, boundary_x, bottom), fill=rgb(COLORS["amber"]), width=3)
            draw.text(
                (boundary_x + 8, top + 8),
                "epoch 2",
                fill=rgb(COLORS["amber"]),
                font=font(16, True),
            )

    loss_path = output_dir / "10_candidate_a_training_loss.png"
    image, draw = chart_canvas(
        "Candidate A training loss — full 348-step run",
        "Raw explicit-negative MNRL loss and a 15-step rolling mean; the two-step smoke is excluded.",
    )
    steps = [float(value) for value in facts["loss_steps"]]
    raw_losses = [float(value) for value in facts["loss_values"]]
    rolling_losses = [float(value) for value in facts["loss_rolling_mean_15"]]
    draw_line_panel(
        draw,
        raw_losses,
        box=(170, 270, 1450, 485),
        color=COLORS["rose"],
        maximum=max(raw_losses),
        label="Raw batch loss",
        x_values=steps,
        epoch_boundary=int(facts["epoch_boundary_step"]),
    )
    draw_line_panel(
        draw,
        rolling_losses,
        box=(170, 595, 1450, 770),
        color=COLORS["sage"],
        maximum=max(rolling_losses),
        label="15-step rolling mean",
        x_values=steps,
        epoch_boundary=int(facts["epoch_boundary_step"]),
        show_x_ticks=True,
    )
    draw.text((735, 790), "optimizer step", fill=rgb(COLORS["muted"]), font=font(18, True))
    image.save(loss_path, dpi=(180, 180))

    dev_path = output_dir / "11_dev_retrieval_comparison.png"
    image, draw = chart_canvas(
        "Held-out development retrieval — base vs Candidate A",
        "18 development queries, MemexAI query prefix, identical 2,018-chunk corpus; y-axis starts at zero.",
    )
    categories = [
        ("Exact MRR", "exact_chunk_mrr"),
        ("Local-window MRR", "local_window_mrr"),
        ("Same-video MRR", "same_video_mrr"),
    ]
    left, top, right, bottom = 180, 300, 1450, 720
    for tick in range(6):
        value = tick / 5
        y = bottom - int(value * (bottom - top))
        draw.line((left, y, right, y), fill=rgb(COLORS["grid"]), width=2)
        draw.text((105, y - 12), f"{value:.1f}", fill=rgb(COLORS["muted"]), font=font(18))
    series = [
        ("Base Qwen 0.6B", facts["base_dev"], COLORS["purple"]),
        ("Candidate A", facts["candidate_dev"], COLORS["sage"]),
    ]
    for legend_index, (label, _, color) in enumerate(series):
        x = 450 + legend_index * 360
        draw.rounded_rectangle((x, 220, x + 28, 248), radius=5, fill=rgb(color))
        draw.text((x + 40, 218), label, fill=rgb(COLORS["ink"]), font=font(20, True))
    group_width = (right - left) / len(categories)
    for category_index, (label, key) in enumerate(categories):
        center = left + group_width * (category_index + 0.5)
        for series_index, (_, metrics, color) in enumerate(series):
            value = float(metrics[key])
            x1 = int(center - 80 + series_index * 90)
            x2 = x1 + 70
            y = bottom - int(value * (bottom - top))
            draw.rounded_rectangle((x1, y, x2, bottom), radius=8, fill=rgb(color))
            draw.text(
                (x1 + 3, y - 27), f"{value:.3f}", fill=rgb(COLORS["ink"]), font=font(17, True)
            )
        bbox = draw.textbbox((0, 0), label, font=font(19, True))
        draw.text(
            (center - (bbox[2] - bbox[0]) / 2, bottom + 22),
            label,
            fill=rgb(COLORS["ink"]),
            font=font(19, True),
        )
    image.save(dev_path, dpi=(180, 180))

    resource_path = output_dir / "12_candidate_a_resource_trace.png"
    image, draw = chart_canvas(
        "Candidate A Ponyo resource trace",
        "Two-second samples across training, checkpointing, and final export; panels use independent labeled scales.",
    )
    minutes = [float(value) for value in facts["resource_minutes"]]
    available = [float(value) for value in facts["available_memory_gib"]]
    gpu = [float(value) for value in facts["gpu_utilization_pct"]]
    draw_line_panel(
        draw,
        available,
        box=(170, 270, 1450, 485),
        color=COLORS["purple"],
        maximum=max(50.0, max(available)),
        label="Available unified memory (GiB)",
        x_values=minutes,
    )
    draw_line_panel(
        draw,
        gpu,
        box=(170, 595, 1450, 770),
        color=COLORS["sage"],
        maximum=100.0,
        label="GPU utilization (%)",
        x_values=minutes,
        show_x_ticks=True,
    )
    draw.text((745, 790), "elapsed minutes", fill=rgb(COLORS["muted"]), font=font(18, True))
    image.save(resource_path, dpi=(180, 180))

    return {
        "coverage": coverage_path,
        "sources": source_path,
        "memory": memory_path,
        "loss": loss_path,
        "dev": dev_path,
        "resources": resource_path,
    }


def write_csv_table(path: Path, facts: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle)
        writer.writerow(["metric", "value"])
        for key in (
            "rows",
            "ready_rows",
            "insufficient_rows",
            "coverage",
            "selected_negatives",
            "excluded_candidates",
            "grouped_positive_rows",
        ):
            writer.writerow([key, facts[key]])


def write_training_csv(path: Path, facts: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle)
        writer.writerow(["step", "epoch", "loss", "rolling_mean_15", "learning_rate"])
        batches_per_epoch = int(facts["epoch_boundary_step"])
        for step, loss, rolling, learning_rate in zip(
            facts["loss_steps"],
            facts["loss_values"],
            facts["loss_rolling_mean_15"],
            facts["learning_rates"],
            strict=True,
        ):
            writer.writerow([step, float(step) / batches_per_epoch, loss, rolling, learning_rate])


def write_dev_comparison_csv(path: Path, facts: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    rows = [
        ("exact_mrr", "exact_chunk_mrr"),
        ("local_window_mrr", "local_window_mrr"),
        ("same_video_mrr", "same_video_mrr"),
    ]
    with path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.writer(handle)
        writer.writerow(
            [
                "metric",
                "base",
                "candidate_a",
                "delta",
                "paired_bootstrap_ci95_low",
                "paired_bootstrap_ci95_high",
                "dev_queries",
            ]
        )
        for label, manifest_key in rows:
            bootstrap = facts["dev_bootstrap"][label]
            writer.writerow(
                [
                    label,
                    facts["base_dev"][manifest_key],
                    facts["candidate_dev"][manifest_key],
                    bootstrap["delta"],
                    bootstrap["ci95_low"],
                    bootstrap["ci95_high"],
                    bootstrap["rows"],
                ]
            )


def write_analysis_summary(path: Path, facts: dict[str, Any]) -> None:
    summary = {
        "promotion_status": facts["promotion_status"],
        "frozen_test_status": facts["frozen_test_status"],
        "training": {
            "result_metrics": facts["training_result"],
            "loss_overall": facts["loss_overall"],
            "loss_epoch_1": facts["loss_epoch_1"],
            "loss_epoch_2": facts["loss_epoch_2"],
        },
        "resources": {
            "sample_count": facts["resource_sample_count"],
            "available_memory_floor_gib": facts["available_memory_floor_gib"],
            "swap_growth_gib": facts["swap_growth_gib"],
            "gpu_utilization_mean_pct": facts["gpu_utilization_mean_pct"],
            "temperature_peak_c": facts["temperature_peak_c"],
        },
        "development": {
            "base": facts["base_dev"],
            "candidate_a": facts["candidate_dev"],
            "paired_bootstrap": facts["dev_bootstrap"],
        },
        "product_fixture": {
            "passed": facts["product_eval"]["passed"],
            "metrics": facts["product_eval"]["metrics"],
        },
        "analysis_decisions": [
            "METRIC: MemexAI-prefixed exact, local-window, and same-video MRR are the pre-specified dev comparison.",
            "UNCERTAINTY: Paired 10,000-sample bootstrap intervals are shown because dev has only 18 queries.",
            "VIZ: Raw and 15-step rolling-mean loss use stacked panels instead of a dual y-axis.",
            "GATE: Frozen 36-query test remains unopened after the development gate failed.",
            "LIMITATION: Product fixture is saturated at MRR/Recall 1.0 and is compatibility evidence, not selection evidence.",
        ],
    }
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(summary, indent=2, sort_keys=True) + "\n", encoding="utf-8")


def write_markdown(path: Path, facts: dict[str, Any], charts: dict[str, Path]) -> None:
    relative = {key: value.relative_to(path.parent).as_posix() for key, value in charts.items()}
    exact_ci = facts["dev_bootstrap"]["exact_mrr"]
    local_ci = facts["dev_bootstrap"]["local_window_mrr"]
    video_ci = facts["dev_bootstrap"]["same_video_mrr"]
    path.write_text(
        f"""# MemexAI Positive-Aware Embedding Campaign

## Decision

**Do not promote Candidate A.** The proper run completed all 348 optimizer steps over 696 positive-aware rows, but it missed the pre-specified held-out development gates. The frozen 36-query test remains unopened because development failed. The SearchTube product fixture still passes, but that fixture is saturated and cannot select between these models.

## Full training run

- Model: Qwen3-Embedding-0.6B, 768-dimensional evaluation, BF16.
- Data: 696 ready rows, exactly three explicit negatives per query, collision-safe by positive video.
- Schedule: two epochs, batch 4, 35 warmup steps, cosine decay, LR 2e-6.
- Runtime: {facts['training_result']['train_runtime'] / 60:.1f} minutes; {facts['training_result']['train_steps_per_second']:.2f} steps/second.
- Mean train loss: {facts['training_result']['train_loss']:.6f}; median per-batch loss: {facts['loss_overall']['median']:.4f}.
- Epoch mean loss: {facts['loss_epoch_1']['mean']:.4f} in epoch 1 and {facts['loss_epoch_2']['mean']:.4f} in epoch 2.
- Checkpoints: quarter-run checkpoints with only the newest two retained; final export matches checkpoint 348 by SHA-256.

![Full Candidate A loss curve]({relative['loss']})

The raw curve is intentionally retained because explicit hard-negative batches are heterogeneous. Twelve epoch-1 batches exceeded loss 1.0; none did in epoch 2. The lower rolling-mean panel improves trend readability without hiding the spikes.

## Held-out development result

The comparison uses 18 held-out development queries, the MemexAI query prefix, each model's own 768d corpus embeddings, and the identical 2,018-chunk corpus.

| Metric | Base Qwen | Candidate A | Delta | Paired bootstrap 95% CI |
|---|---:|---:|---:|---:|
| Exact MRR | {facts['base_dev']['exact_chunk_mrr']:.6f} | {facts['candidate_dev']['exact_chunk_mrr']:.6f} | {exact_ci['delta']:+.6f} | [{exact_ci['ci95_low']:+.6f}, {exact_ci['ci95_high']:+.6f}] |
| Local-window MRR | {facts['base_dev']['local_window_mrr']:.6f} | {facts['candidate_dev']['local_window_mrr']:.6f} | {local_ci['delta']:+.6f} | [{local_ci['ci95_low']:+.6f}, {local_ci['ci95_high']:+.6f}] |
| Same-video MRR | {facts['base_dev']['same_video_mrr']:.6f} | {facts['candidate_dev']['same_video_mrr']:.6f} | {video_ci['delta']:+.6f} | [{video_ci['ci95_low']:+.6f}, {video_ci['ci95_high']:+.6f}] |

![Held-out development comparison]({relative['dev']})

The intervals are wide because `n=18`, so this run does not prove the recipe is universally harmful. It does fail the campaign's pre-specified point gates, which is sufficient to block promotion and protect the frozen test from tuning leakage.

## Product compatibility

The SearchTube fixture passes: MRR and Recall@1 remain 1.0, wrong-top and hard-negative-above-positive rates remain 0, and nDCG@5 is {facts['product_eval']['metrics']['ndcgAt5']:.6f}. Because the base model already scores essentially perfectly, this is a compatibility check rather than evidence of improvement.

## Dataset findings

- {facts['ready_rows']} of {facts['rows']} queries ({facts['coverage']:.1%}) have at least three safe explicit negatives.
- {facts['selected_negatives']:,} negatives were selected: {facts['negative_sources'].get('bm25', 0):,} BM25-mined and {facts['negative_sources'].get('labeled', 0):,} previously labeled.
- {facts['excluded_candidates']:,} unsafe candidates were excluded, dominated by unreviewed same-video and adjacent windows.
- {facts['grouped_positive_rows']} rows carry multiple labeled positives; the remaining rows keep the exact positive only until reviewed.
- {facts['insufficient_rows']} rows are explicitly marked insufficient and omitted from explicit-negative training.

![Positive-aware coverage]({relative['coverage']})

![Negative sources and exclusions]({relative['sources']})

![Expanded batch memory guard]({relative['memory']})

## Resource envelope

The corrected failed batch-16 probe contained 80 encoder sequences, not 144: the legacy Sentence Transformers path silently truncated variable negative lanes to three. The direct trainer now enforces exactly three lanes, uses the collision-safe sampler at runtime, and rejects non-cached batches above 64 sequences before loading model weights. Candidate A therefore ran at 20 sequences per batch.

![Ponyo resource trace]({relative['resources']})

Across {facts['resource_sample_count']} two-second samples, available unified memory reached a {facts['available_memory_floor_gib']:.2f} GiB floor, swap grew by {facts['swap_growth_gib']:.2f} GiB, mean GPU utilization was {facts['gpu_utilization_mean_pct']:.1f}%, and peak temperature was {facts['temperature_peak_c']:.0f} C. Batch 4 remains the safe explicit-negative envelope while the resident services share unified memory.

## Interpretation and next experiment

The training curve shows rapid memorization of the selected hard-negative objective, while held-out retrieval worsens. Plausible explanations include overly aggressive explicit-negative pressure, BM25-heavy mining that does not match dense decision boundaries, too few reviewed equivalent positives, and loss scale/temperature that separates training groups without preserving global corpus geometry. These are hypotheses, not causal findings.

1. Keep base Qwen as champion; do not release Candidate A.
2. Add base-Qwen dense top-50 candidates and true RRF fusion before rebuilding negatives.
3. Complete the 30-query double-label pilot for equivalent positives and false-negative adjudication.
4. Run a smaller-step/lower-pressure ablation (one epoch and/or higher temperature) and evaluate dev checkpoints before a full second epoch.
5. Test CachedMNRL only after sampler/semantics and memory proofs; do not assume a larger effective batch fixes negative quality.
6. Open a new frozen test only after a future candidate passes development gates.

## Repository boundary

BashGym trains, evaluates, visualizes, and packages the model. MemexAI consumes a released model and integration manifest. Any fusion or reranker improvement is implemented and rolled out separately in the MemexAI repository; there is no repository merge or BashGym runtime dependency.

## Analysis decisions and limitations

- Loss charts exclude the two-step smoke and use all 348 full-run observations.
- MRR deltas use a deterministic 10,000-sample paired bootstrap because development has only 18 queries.
- No multiple slice claims are promoted from this small set.
- The frozen test was not opened after the failed development gate.
- BashGym remains the independent training/evaluation system; no repository merge or runtime dependency was introduced.
""",
        encoding="utf-8",
    )


def extend_docx(
    source_docx: Path,
    output_docx: Path,
    facts: dict[str, Any],
    charts: dict[str, Path],
) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches

    from scripts.memexai.build_embedding_experiment_report import (
        COLORS,
        add_body,
        add_callout,
        add_figure,
        add_heading,
        add_table,
        set_run_font,
    )

    doc = Document(source_docx)
    doc.add_page_break()
    add_heading(doc, "Approach B campaign update", 1)
    add_body(
        doc,
        "BashGym built deterministic grouped-positive training records, ran a complete positive-aware candidate, and evaluated it against the frozen development contract. This update preserves the original report and adds measured training, resource, compatibility, and held-out development evidence.",
    )
    add_callout(
        doc,
        "Decision: do not promote Candidate A",
        "The 348-step BF16 run completed successfully, but Candidate A regressed exact, local-window, and same-video MRR on the 18-query held-out development set. The frozen 36-query test remains unopened. Base Qwen stays champion.",
        fill=COLORS["amber_light"],
    )
    add_figure(
        doc,
        charts["coverage"],
        "Figure 7. Positive-aware explicit-negative coverage across Real702.",
    )
    add_table(
        doc,
        ["Measure", "Value", "Meaning"],
        [
            ["Ready rows", str(facts["ready_rows"]), "At least 3 safe negatives"],
            [
                "Insufficient rows",
                str(facts["insufficient_rows"]),
                "Excluded from explicit-negative training",
            ],
            [
                "Selected negatives",
                f"{facts['selected_negatives']:,}",
                "Labeled plus BM25 candidates",
            ],
            [
                "Grouped-positive rows",
                str(facts["grouped_positive_rows"]),
                "More than one labeled positive",
            ],
        ],
        [2600, 1500, 5260],
    )
    doc.add_page_break()
    add_heading(doc, "False-negative protection", 1)
    add_figure(doc, charts["sources"], "Figure 8. Selected negative lanes and exclusion reasons.")
    add_body(
        doc,
        "The builder excludes exact/equivalent positives, adjacent or overlapping same-video windows, near-duplicate text, and every unreviewed same-video candidate. This intentionally sacrifices some hard negatives to avoid repeating the false-negative failure mode observed in Real702.",
    )
    add_callout(
        doc,
        "Repository boundary",
        "BashGym owns model development and reporting. MemexAI may separately implement a benchmark-supported reranker or fusion change, consuming only released model and integration artifacts.",
        fill=COLORS["purple_light"],
    )
    add_figure(
        doc,
        charts["memory"],
        "Figure 9. Expanded-sequence memory guard derived from the Ponyo batch probe.",
    )
    doc.add_page_break()
    add_heading(doc, "Full Candidate A training run", 1)
    add_body(
        doc,
        f"Candidate A trained on 696 rows for 348 optimizer steps over two epochs in BF16. The direct trainer enforced exactly three explicit negatives and one positive-video collision group per batch. Trainer runtime was {facts['training_result']['train_runtime'] / 60:.1f} minutes at {facts['training_result']['train_steps_per_second']:.2f} steps per second.",
    )
    add_figure(
        doc,
        charts["loss"],
        "Figure 10. Full-run raw batch loss and 15-step rolling mean; smoke steps are excluded.",
    )
    add_table(
        doc,
        ["Loss summary", "Epoch 1", "Epoch 2"],
        [
            [
                "Mean",
                f"{facts['loss_epoch_1']['mean']:.4f}",
                f"{facts['loss_epoch_2']['mean']:.4f}",
            ],
            [
                "Median",
                f"{facts['loss_epoch_1']['median']:.4f}",
                f"{facts['loss_epoch_2']['median']:.4f}",
            ],
            [
                "90th percentile",
                f"{facts['loss_epoch_1']['p90']:.4f}",
                f"{facts['loss_epoch_2']['p90']:.4f}",
            ],
            [
                "Batches above 1.0",
                str(facts["loss_epoch_1"]["above_1_count"]),
                str(facts["loss_epoch_2"]["above_1_count"]),
            ],
        ],
        [3100, 3130, 3130],
    )
    doc.add_page_break()
    add_heading(doc, "Held-out development gate", 1)
    add_figure(
        doc,
        charts["dev"],
        "Figure 11. Base Qwen versus Candidate A on 18 held-out development queries.",
    )
    add_table(
        doc,
        ["Metric", "Base", "Candidate A", "Delta"],
        [
            [
                "Exact MRR",
                f"{facts['base_dev']['exact_chunk_mrr']:.4f}",
                f"{facts['candidate_dev']['exact_chunk_mrr']:.4f}",
                f"{facts['dev_bootstrap']['exact_mrr']['delta']:+.4f}",
            ],
            [
                "Local-window MRR",
                f"{facts['base_dev']['local_window_mrr']:.4f}",
                f"{facts['candidate_dev']['local_window_mrr']:.4f}",
                f"{facts['dev_bootstrap']['local_window_mrr']['delta']:+.4f}",
            ],
            [
                "Same-video MRR",
                f"{facts['base_dev']['same_video_mrr']:.4f}",
                f"{facts['candidate_dev']['same_video_mrr']:.4f}",
                f"{facts['dev_bootstrap']['same_video_mrr']['delta']:+.4f}",
            ],
        ],
        [3100, 2000, 2200, 2060],
    )
    add_body(
        doc,
        "Paired 10,000-sample bootstrap intervals are wide because development contains only 18 queries. The result does not prove the recipe is universally harmful, but it misses the pre-specified point gates and therefore blocks promotion without opening the frozen test.",
    )
    add_callout(
        doc,
        "Product fixture",
        f"Compatibility remains intact: the SearchTube fixture passed with MRR and Recall@1 at 1.0 and nDCG@5 at {facts['product_eval']['metrics']['ndcgAt5']:.6f}. The fixture is saturated, so it cannot select the champion.",
        fill=COLORS["sage_light"],
    )
    doc.add_page_break()
    add_heading(doc, "Ponyo resource envelope and next experiment", 1)
    add_figure(
        doc,
        charts["resources"],
        "Figure 12. Unified-memory availability and GPU utilization across the full run.",
    )
    add_body(
        doc,
        f"Available unified memory reached {facts['available_memory_floor_gib']:.2f} GiB, swap grew {facts['swap_growth_gib']:.2f} GiB, mean GPU utilization was {facts['gpu_utilization_mean_pct']:.1f}%, and peak temperature was {facts['temperature_peak_c']:.0f} C. Batch 4 remains the safe explicit-negative envelope while Ponyo's resident services share unified memory.",
    )
    add_callout(
        doc,
        "Recommended next iteration",
        "Keep base Qwen as champion. Improve negative quality with base-Qwen dense top-50 mining plus RRF, complete the 30-query double-label pilot, then run a lower-pressure one-epoch/checkpoint ablation. CachedMNRL remains evidence-gated; a larger effective batch cannot repair mislabeled negatives.",
        fill=COLORS["purple_light"],
    )
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    for section in doc.sections:
        for footer in (
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            paragraph = footer.paragraphs[0]
            paragraph.clear()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.right_indent = Inches(0.15)
            label_run = paragraph.add_run("Page ")
            set_run_font(label_run, size=9, color=COLORS["muted"])
            field_run = paragraph.add_run()
            set_run_font(field_run, size=9, color=COLORS["muted"])
            field_begin = OxmlElement("w:fldChar")
            field_begin.set(qn("w:fldCharType"), "begin")
            instruction = OxmlElement("w:instrText")
            instruction.set(qn("xml:space"), "preserve")
            instruction.text = " PAGE "
            field_separate = OxmlElement("w:fldChar")
            field_separate.set(qn("w:fldCharType"), "separate")
            display = OxmlElement("w:t")
            display.text = "1"
            field_end = OxmlElement("w:fldChar")
            field_end.set(qn("w:fldCharType"), "end")
            for element in (field_begin, instruction, field_separate, display, field_end):
                field_run._r.append(element)
    doc.core_properties.title = "MemexAI Embedding Experiments and Positive-Aware Campaign"
    doc.save(output_docx)


def write_pdf(path: Path, facts: dict[str, Any], charts: dict[str, Path]) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        Image,
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CampaignTitle",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=25,
            leading=29,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CampaignCaption",
            parent=styles["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=8.5,
            leading=11,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#667085"),
            spaceAfter=10,
        )
    )
    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=0.8 * inch,
        rightMargin=0.8 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        title="MemexAI Positive-Aware Embedding Campaign",
        author="BashGym / MemexAI",
    )
    story: list[Any] = [
        Spacer(1, 0.55 * inch),
        Paragraph("TECHNICAL CAMPAIGN UPDATE", styles["Heading3"]),
        Paragraph("MemexAI Positive-Aware<br/>Embedding Campaign", styles["CampaignTitle"]),
        Paragraph(
            "Grouped positives, safe hard-negative mining, explicit-negative training, and the independent BashGym-to-MemexAI handoff.",
            styles["Heading2"],
        ),
        Spacer(1, 0.2 * inch),
        Paragraph(
            "Decision: do not promote Candidate A. The complete 348-step run succeeded operationally but regressed the pre-specified held-out development metrics. Base Qwen remains champion and the frozen test remains unopened.",
            styles["BodyText"],
        ),
        PageBreak(),
        Paragraph("Measured dataset findings", styles["Heading1"]),
        Paragraph(
            f"{facts['ready_rows']} of {facts['rows']} queries ({facts['coverage']:.1%}) have at least three safe explicit negatives. {facts['insufficient_rows']} rows are explicitly excluded from this training arm rather than silently weakened.",
            styles["BodyText"],
        ),
        Image(str(charts["coverage"]), width=6.5 * inch, height=3.66 * inch),
        Paragraph("Figure 1. Positive-aware coverage across Real702.", styles["CampaignCaption"]),
    ]
    data = [
        ["Measure", "Value", "Interpretation"],
        ["Ready rows", str(facts["ready_rows"]), "At least three safe negatives"],
        ["Selected negatives", f"{facts['selected_negatives']:,}", "Labeled plus BM25-mined"],
        [
            "Excluded candidates",
            f"{facts['excluded_candidates']:,}",
            "Potential false negatives removed",
        ],
        [
            "Grouped-positive rows",
            str(facts["grouped_positive_rows"]),
            "Multiple labeled positives",
        ],
    ]
    table = Table(data, colWidths=[1.55 * inch, 1.05 * inch, 3.9 * inch], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#1F2937")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D9DEE7")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend([table, PageBreak(), Paragraph("False-negative protection", styles["Heading1"])])
    story.extend(
        [
            Image(str(charts["sources"]), width=6.5 * inch, height=3.66 * inch),
            Paragraph(
                "Figure 2. Selected negative lanes and exclusion reasons.",
                styles["CampaignCaption"],
            ),
            Paragraph(
                "The builder excludes exact/equivalent positives, adjacent or overlapping same-video windows, near-duplicate text, and every unreviewed same-video candidate. This conservative policy directly addresses the false-negative failure mode inferred from Real702.",
                styles["BodyText"],
            ),
            Spacer(1, 0.25 * inch),
            Paragraph("Launch and repository boundary", styles["Heading1"]),
            Paragraph(
                "The corrected batch-16 probe contained 80 encoder sequences because the legacy trainer silently truncated variable negative lanes to three. The direct trainer now enforces exactly three lanes, verifies the collision-safe sampler at runtime, and refuses non-cached batches above 64 sequences. Candidate A ran at physical batch 4, or 20 encoder sequences.",
                styles["BodyText"],
            ),
            Image(str(charts["memory"]), width=6.5 * inch, height=3.66 * inch),
            Paragraph(
                "Figure 3. The memory guard converts the failed Ponyo probe into a launch invariant.",
                styles["CampaignCaption"],
            ),
            Spacer(1, 0.15 * inch),
            Paragraph(
                "BashGym owns data design, training, evaluation, visualization, and packaging. MemexAI consumes only a released model plus immutable integration manifest. Any fusion or reranker change is implemented independently in the MemexAI repository with no BashGym runtime dependency and no repository merge.",
                styles["BodyText"],
            ),
            Spacer(1, 0.2 * inch),
            Paragraph("Next evidence", styles["Heading2"]),
            Paragraph(
                "1. Keep base Qwen as champion.<br/>2. Add base-Qwen dense top-50 mining and RRF fusion.<br/>3. Run the 30-query double-label pilot.<br/>4. Test a lower-pressure one-epoch/checkpoint ablation on development only.<br/>5. Keep the frozen test closed until a future candidate passes development.",
                styles["BodyText"],
            ),
        ]
    )
    dev_data = [
        ["Metric", "Base", "Candidate A", "Delta"],
        [
            "Exact MRR",
            f"{facts['base_dev']['exact_chunk_mrr']:.4f}",
            f"{facts['candidate_dev']['exact_chunk_mrr']:.4f}",
            f"{facts['dev_bootstrap']['exact_mrr']['delta']:+.4f}",
        ],
        [
            "Local-window MRR",
            f"{facts['base_dev']['local_window_mrr']:.4f}",
            f"{facts['candidate_dev']['local_window_mrr']:.4f}",
            f"{facts['dev_bootstrap']['local_window_mrr']['delta']:+.4f}",
        ],
        [
            "Same-video MRR",
            f"{facts['base_dev']['same_video_mrr']:.4f}",
            f"{facts['candidate_dev']['same_video_mrr']:.4f}",
            f"{facts['dev_bootstrap']['same_video_mrr']['delta']:+.4f}",
        ],
    ]
    dev_table = Table(
        dev_data, colWidths=[2.3 * inch, 1.2 * inch, 1.4 * inch, 1.1 * inch], repeatRows=1
    )
    dev_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D9DEE7")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend(
        [
            PageBreak(),
            Paragraph("Full Candidate A training run", styles["Heading1"]),
            Paragraph(
                f"Candidate A completed 348 BF16 optimizer steps over 696 rows in {facts['training_result']['train_runtime'] / 60:.1f} minutes. Mean train loss was {facts['training_result']['train_loss']:.6f}; epoch mean loss moved from {facts['loss_epoch_1']['mean']:.4f} to {facts['loss_epoch_2']['mean']:.4f}.",
                styles["BodyText"],
            ),
            Image(str(charts["loss"]), width=6.5 * inch, height=3.66 * inch),
            Paragraph(
                "Figure 4. Raw full-run loss and 15-step rolling mean; smoke steps are excluded.",
                styles["CampaignCaption"],
            ),
            Paragraph(
                "The rapid loss reduction shows the selected training objective was learned. It does not establish retrieval improvement; the held-out comparison below is the promotion evidence.",
                styles["BodyText"],
            ),
            PageBreak(),
            Paragraph("Held-out development gate", styles["Heading1"]),
            Image(str(charts["dev"]), width=6.5 * inch, height=3.66 * inch),
            Paragraph(
                "Figure 5. Base Qwen versus Candidate A on 18 held-out development queries.",
                styles["CampaignCaption"],
            ),
            dev_table,
            Spacer(1, 0.18 * inch),
            Paragraph(
                "Candidate A missed all three directional MRR gates. Paired 10,000-sample bootstrap intervals remain wide at n=18, so this is a campaign gate failure rather than a universal causal claim. The frozen 36-query test was not opened.",
                styles["BodyText"],
            ),
            Paragraph(
                f"The separate SearchTube fixture passed with MRR/Recall@1 at 1.0 and nDCG@5 at {facts['product_eval']['metrics']['ndcgAt5']:.6f}; it is saturated and provides compatibility, not selection, evidence.",
                styles["BodyText"],
            ),
            PageBreak(),
            Paragraph("Ponyo resource envelope", styles["Heading1"]),
            Image(str(charts["resources"]), width=6.5 * inch, height=3.66 * inch),
            Paragraph(
                "Figure 6. Unified-memory availability and GPU utilization across the full run.",
                styles["CampaignCaption"],
            ),
            Paragraph(
                f"Available unified memory reached {facts['available_memory_floor_gib']:.2f} GiB, swap grew {facts['swap_growth_gib']:.2f} GiB, mean GPU utilization was {facts['gpu_utilization_mean_pct']:.1f}%, and peak temperature was {facts['temperature_peak_c']:.0f} C. Batch 4 remains the safe explicit-negative envelope with resident services active.",
                styles["BodyText"],
            ),
            Spacer(1, 0.16 * inch),
            Paragraph("Recommended next iteration", styles["Heading2"]),
            Paragraph(
                "Keep base Qwen as champion. Improve negative quality with base-Qwen dense top-50 mining plus RRF, complete the 30-query double-label pilot, and test a lower-pressure one-epoch/checkpoint ablation. CachedMNRL stays gated on semantics and memory proof; a larger batch cannot repair false negatives.",
                styles["BodyText"],
            ),
        ]
    )
    doc.build(story)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--positive-manifest", type=Path, required=True)
    parser.add_argument("--source-docx", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--training-manifest", type=Path, required=True)
    parser.add_argument("--training-metrics-jsonl", type=Path, required=True)
    parser.add_argument("--resource-samples-csv", type=Path, required=True)
    parser.add_argument("--base-dev-manifest", type=Path, required=True)
    parser.add_argument("--candidate-dev-manifest", type=Path, required=True)
    parser.add_argument("--base-dev-rows", type=Path, required=True)
    parser.add_argument("--candidate-dev-rows", type=Path, required=True)
    parser.add_argument("--product-eval-json", type=Path, required=True)
    args = parser.parse_args()
    manifest = load_json(args.positive_manifest)
    facts = report_facts(manifest)
    add_run_facts(
        facts,
        training_manifest_path=args.training_manifest,
        training_metrics_path=args.training_metrics_jsonl,
        resource_samples_path=args.resource_samples_csv,
        base_dev_manifest_path=args.base_dev_manifest,
        candidate_dev_manifest_path=args.candidate_dev_manifest,
        base_dev_rows_path=args.base_dev_rows,
        candidate_dev_rows_path=args.candidate_dev_rows,
        product_eval_path=args.product_eval_json,
    )
    charts = build_charts(args.output_dir / "charts", facts)
    write_csv_table(args.output_dir / "tables" / "positive_aware_summary.csv", facts)
    write_training_csv(args.output_dir / "tables" / "candidate_a_training_metrics.csv", facts)
    write_dev_comparison_csv(args.output_dir / "tables" / "candidate_a_dev_comparison.csv", facts)
    write_analysis_summary(args.output_dir / "analysis_summary.json", facts)
    write_markdown(args.output_dir / "MemexAI_Positive_Aware_Campaign_Report.md", facts, charts)
    extend_docx(
        args.source_docx,
        args.output_dir / "MemexAI_Positive_Aware_Campaign_Report.docx",
        facts,
        charts,
    )
    write_pdf(args.output_dir / "MemexAI_Positive_Aware_Campaign_Report.pdf", facts, charts)
    print(args.output_dir.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
