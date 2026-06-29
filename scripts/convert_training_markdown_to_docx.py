"""Convert BashGym training markdown references into styled Word documents."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BODY_FONT = "Calibri"
MONO_FONT = "Consolas"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
TEXT = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x55, 0x65, 0x73)
TABLE_HEADER_FILL = "E8EEF5"
CODE_FILL = "F6F8FA"
BORDER = "D7DEE8"
CONTENT_WIDTH_IN = 6.5


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color: str = BORDER) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.first_child_found_in("w:pBdr")
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def configure_document(doc: Document, title: str, source: Path) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    title_style = styles["Title"]
    title_style.font.name = BODY_FONT
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    title_style.font.size = Pt(20)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title_style.paragraph_format.space_after = Pt(10)

    heading_1 = styles["Heading 1"]
    heading_1.font.name = BODY_FONT
    heading_1._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    heading_1.font.size = Pt(16)
    heading_1.font.bold = True
    heading_1.font.color.rgb = BLUE
    heading_1.paragraph_format.space_before = Pt(18)
    heading_1.paragraph_format.space_after = Pt(10)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = styles["Heading 2"]
    heading_2.font.name = BODY_FONT
    heading_2._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    heading_2.font.size = Pt(13)
    heading_2.font.bold = True
    heading_2.font.color.rgb = BLUE
    heading_2.paragraph_format.space_before = Pt(14)
    heading_2.paragraph_format.space_after = Pt(7)
    heading_2.paragraph_format.keep_with_next = True

    heading_3 = styles["Heading 3"]
    heading_3.font.name = BODY_FONT
    heading_3._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    heading_3.font.size = Pt(12)
    heading_3.font.bold = True
    heading_3.font.color.rgb = DARK_BLUE
    heading_3.paragraph_format.space_before = Pt(10)
    heading_3.paragraph_format.space_after = Pt(5)
    heading_3.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.187)
        style.paragraph_format.space_after = Pt(2)

    if "BashGym Code" not in styles:
        code_style = styles.add_style("BashGym Code", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["BashGym Code"]
    code_style.font.name = MONO_FONT
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), MONO_FONT)
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    code_style.paragraph_format.left_indent = Inches(0.12)
    code_style.paragraph_format.right_indent = Inches(0.08)
    code_style.paragraph_format.space_before = Pt(0)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if "BashGym Quote" not in styles:
        quote_style = styles.add_style("BashGym Quote", WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote_style = styles["BashGym Quote"]
    quote_style.font.name = BODY_FONT
    quote_style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    quote_style.font.size = Pt(11)
    quote_style.font.italic = True
    quote_style.font.color.rgb = MUTED
    quote_style.paragraph_format.left_indent = Inches(0.25)
    quote_style.paragraph_format.space_after = Pt(6)

    footer = section.footer.paragraphs[0]
    footer.text = f"{title} | Source: {source.as_posix()}"
    footer.style = styles["Footer"]
    for run in footer.runs:
        run.font.name = BODY_FONT
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED


def convert_links(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)


def clean_inline_text(text: str) -> str:
    text = convert_links(text)
    text = text.replace("\\_", "_")
    return text


def add_inline_runs(paragraph, text: str, *, base_font_size: float | None = None) -> None:
    text = clean_inline_text(text)
    token = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in token.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            if base_font_size:
                run.font.size = Pt(base_font_size)
        piece = match.group(0)
        if piece.startswith("**"):
            run = paragraph.add_run(piece[2:-2])
            run.bold = True
            if base_font_size:
                run.font.size = Pt(base_font_size)
        elif piece.startswith("`"):
            run = paragraph.add_run(piece[1:-1])
            run.font.name = MONO_FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), MONO_FONT)
            run.font.size = Pt(base_font_size or 9)
            run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if base_font_size:
            run.font.size = Pt(base_font_size)


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    cells = split_table_row(stripped)
    return len(cells) > 1 and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (column_count - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized_rows), cols=column_count)
    table.autofit = False
    table.allow_autofit = False
    set_table_borders(table)

    width = CONTENT_WIDTH_IN / column_count
    font_size = 8.5 if column_count <= 4 else 8

    for row_index, row in enumerate(normalized_rows):
        keep_table_row_together(table.rows[row_index])
        if row_index == 0:
            repeat_table_header(table.rows[row_index])
        for column_index, text in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.width = Inches(width)
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, TABLE_HEADER_FILL)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            add_inline_runs(paragraph, text, base_font_size=font_size)
            for run in paragraph.runs:
                run.font.name = BODY_FONT
                run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
                run.font.size = Pt(font_size)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str], language: str) -> None:
    if language:
        label = doc.add_paragraph(style="BashGym Code")
        add_inline_runs(label, f"# {language}")
        set_paragraph_shading(label, CODE_FILL)
    for line in lines or [""]:
        paragraph = doc.add_paragraph(style="BashGym Code")
        paragraph.paragraph_format.keep_together = False
        set_paragraph_shading(paragraph, CODE_FILL)
        run = paragraph.add_run(line.expandtabs(4))
        run.font.name = MONO_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), MONO_FONT)
        run.font.size = Pt(9)
    doc.add_paragraph()


def add_divider(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    set_paragraph_bottom_border(paragraph)


def add_heading_or_title(doc: Document, text: str, level: int, seen_title: bool) -> bool:
    if level == 1 and not seen_title:
        paragraph = doc.add_paragraph(style="Title")
        add_inline_runs(paragraph, text)
        return True
    style_level = min(max(level, 1), 3)
    paragraph = doc.add_heading(level=style_level)
    add_inline_runs(paragraph, text)
    return seen_title


def add_list_item(doc: Document, text: str, numbered: bool) -> None:
    style = "List Number" if numbered else "List Bullet"
    paragraph = doc.add_paragraph(style=style)
    add_inline_runs(paragraph, text)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    add_inline_runs(paragraph, text)


def parse_markdown(doc: Document, source_text: str) -> None:
    lines = source_text.splitlines()
    index = 0
    seen_title = False
    in_code = False
    code_language = ""
    code_lines: list[str] = []
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_paragraph(doc, " ".join(paragraph_buffer))
            paragraph_buffer = []

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.rstrip()

        if in_code:
            if line.startswith("```"):
                add_code_block(doc, code_lines, code_language)
                code_lines = []
                code_language = ""
                in_code = False
            else:
                code_lines.append(raw_line)
            index += 1
            continue

        if line.startswith("```"):
            flush_paragraph()
            in_code = True
            code_language = line[3:].strip()
            code_lines = []
            index += 1
            continue

        if not line.strip():
            flush_paragraph()
            index += 1
            continue

        if re.fullmatch(r"-{3,}", line.strip()):
            flush_paragraph()
            add_divider(doc)
            index += 1
            continue

        if (
            line.strip().startswith("|")
            and index + 1 < len(lines)
            and is_table_separator(lines[index + 1])
        ):
            flush_paragraph()
            table_rows = [split_table_row(line)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append(split_table_row(lines[index]))
                index += 1
            add_table(doc, table_rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            text = heading.group(2).strip()
            seen_title = add_heading_or_title(doc, text, level, seen_title)
            index += 1
            continue

        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        if bullet:
            flush_paragraph()
            add_list_item(doc, bullet.group(1).strip(), numbered=False)
            index += 1
            continue

        numbered = re.match(r"^\s*\d+[.)]\s+(.+)$", line)
        if numbered:
            flush_paragraph()
            add_list_item(doc, numbered.group(1).strip(), numbered=True)
            index += 1
            continue

        quote = re.match(r"^>\s?(.+)$", line)
        if quote:
            flush_paragraph()
            paragraph = doc.add_paragraph(style="BashGym Quote")
            add_inline_runs(paragraph, quote.group(1).strip())
            index += 1
            continue

        paragraph_buffer.append(line.strip())
        index += 1

    flush_paragraph()
    if in_code:
        add_code_block(doc, code_lines, code_language)


def convert_file(input_path: Path, output_dir: Path) -> Path:
    source_text = input_path.read_text(encoding="utf-8")
    title_match = re.search(r"^#\s+(.+)$", source_text, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else input_path.stem.replace("-", " ").title()

    doc = Document()
    configure_document(doc, title, input_path)
    parse_markdown(doc, source_text)

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"{input_path.stem}.docx"
    doc.core_properties.title = title
    doc.core_properties.subject = "BashGym training documentation"
    doc.core_properties.author = "BashGym"
    doc.save(output_path)
    return output_path


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("inputs", nargs="+", type=Path)
    parser.add_argument("--output-dir", type=Path, default=Path("docs/training"))
    args = parser.parse_args()

    for input_path in args.inputs:
        output_path = convert_file(input_path, args.output_dir)
        print(output_path)


if __name__ == "__main__":
    main()
