"""Portable campaign charts and publication documents.

The renderers are deliberately projections of the already-sanitized campaign
export snapshot.  They never read training directories or remote logs directly.
"""

from __future__ import annotations

import io
import zipfile
from datetime import UTC, datetime
from pathlib import Path
from typing import Any


class CampaignReportingUnavailableError(RuntimeError):
    """Raised when the optional reporting runtime is not installed."""


_FIXED_TIME = datetime(2000, 1, 1, tzinfo=UTC)
_INK = "332F2A"
_ACCENT = "7C3AED"
_MUTED = "6B645D"
_PAPER = "FFFDF7"
_TABLE_FILL = "F2EEF9"


def _series(snapshot: dict[str, Any]) -> list[tuple[str, str, list[tuple[int, float]]]]:
    by_id = {
        item.get("attempt_id"): item for item in snapshot["attempts"] if item.get("attempt_id")
    }
    result = []
    for attempt_id, values in snapshot["loss_by_attempt"].items():
        points = [
            (int(item["step"]), float(item["value"]))
            for item in values
            if "step" in item and "value" in item
        ]
        if points:
            result.append((attempt_id, by_id.get(attempt_id, {}).get("stage", "unknown"), points))
    return result


def write_loss_png(snapshot: dict[str, Any], path: Path) -> None:
    """Render the deterministic loss projection as a portable PNG."""

    try:
        from PIL import Image, ImageDraw, ImageFont
    except ImportError as exc:  # pragma: no cover - installation-dependent
        raise CampaignReportingUnavailableError("Install bashgym[reports] for PNG exports") from exc

    width, height = 960, 480
    left, right, top, bottom = 70, 30, 50, 65
    image = Image.new("RGB", (width, height), "#fffdf7")
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default()
    draw.text((left, 22), "CAMPAIGN TRAINING LOSS", fill="#332f2a", font=font)
    draw.line((left, height - bottom, width - right, height - bottom), fill="#332f2a", width=2)
    draw.line((left, top, left, height - bottom), fill="#332f2a", width=2)
    draw.text((width // 2 - 35, height - 32), "Training step", fill="#332f2a", font=font)
    draw.text((8, top), "Loss", fill="#332f2a", font=font)
    series = _series(snapshot)
    max_step = max((step for _id, _stage, points in series for step, _ in points), default=1)
    losses = [value for _id, _stage, points in series for _, value in points]
    low, high = min(losses, default=0.0), max(losses, default=1.0)
    if high == low:
        high = low + 1.0

    def xy(step: int, value: float) -> tuple[int, int]:
        x = left + round((width - left - right) * step / max_step)
        y = top + round((height - top - bottom) * (high - value) / (high - low))
        return x, y

    colors = ("#7c3aed", "#d97706", "#15803d", "#0369a1", "#be123c")
    for index, (attempt_id, stage, points) in enumerate(series):
        color = colors[index % len(colors)]
        coordinates = [xy(step, value) for step, value in points]
        if stage == "smoke_training":
            for start, end in zip(coordinates, coordinates[1:], strict=False):
                segments = 12
                for segment in range(0, segments, 2):
                    a = segment / segments
                    b = min((segment + 1) / segments, 1.0)
                    draw.line(
                        (
                            round(start[0] + (end[0] - start[0]) * a),
                            round(start[1] + (end[1] - start[1]) * a),
                            round(start[0] + (end[0] - start[0]) * b),
                            round(start[1] + (end[1] - start[1]) * b),
                        ),
                        fill=color,
                        width=3,
                    )
        elif len(coordinates) > 1:
            draw.line(coordinates, fill=color, width=3)
        elif coordinates:
            x, y = coordinates[0]
            draw.ellipse((x - 2, y - 2, x + 2, y + 2), fill=color)
        label = f"{attempt_id} | {stage}"
        draw.text((left + 10, top + 12 + index * 16), label, fill=color, font=font)
    if not series:
        draw.text((width // 2 - 65, height // 2), "No persisted loss series", fill="#6b645d", font=font)
    draw.text((width - 224, height - 18), "Dashed = smoke engineering evidence", fill="#6b645d", font=font)
    image.save(path, format="PNG", optimize=False, compress_level=9)


def _quality_ready(snapshot: dict[str, Any]) -> bool:
    return bool(
        any(
            item.get("stage") == "full_training" and item.get("status") == "completed"
            for item in snapshot["attempts"]
        )
        and snapshot["comparisons"]
    )


def _quality_text(snapshot: dict[str, Any]) -> str:
    if not _quality_ready(snapshot):
        return (
            "No model-quality findings are claimed. A completed full-training attempt and "
            "persisted comparison are both required."
        )
    latest = snapshot["comparisons"][-1]
    return (
        f"The latest deterministic development gate verdict is "
        f"{latest.get('verdict', 'unknown')}. This finding is backed by a completed "
        "full-training attempt and a persisted comparison."
    )


def _normalize_docx(path: Path) -> None:
    """Remove ZIP wall-clock timestamps so identical evidence produces identical DOCX bytes."""

    with zipfile.ZipFile(path, "r") as source:
        members = [(item.filename, source.read(item.filename)) for item in source.infolist()]
    payload = io.BytesIO()
    with zipfile.ZipFile(payload, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as target:
        for name, content in sorted(members):
            info = zipfile.ZipInfo(name, (2000, 1, 1, 0, 0, 0))
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            target.writestr(info, content)
    path.write_bytes(payload.getvalue())


def write_campaign_docx(
    snapshot: dict[str, Any], source_digest: str, loss_png: Path, path: Path
) -> None:
    """Create an editable standard-business-brief Word report."""

    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:  # pragma: no cover - installation-dependent
        raise CampaignReportingUnavailableError("Install bashgym[reports] for DOCX exports") from exc

    def set_font(run, size: float, color: str = _INK, *, bold: bool = False) -> None:
        run.font.name = "Calibri"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold

    def shade(cell, fill: str) -> None:
        properties = cell._tc.get_or_add_tcPr()
        node = properties.find(qn("w:shd"))
        if node is None:
            node = OxmlElement("w:shd")
            properties.append(node)
        node.set(qn("w:fill"), fill)

    def set_table_geometry(table, widths: tuple[int, ...]) -> None:
        table.autofit = False
        properties = table._tbl.tblPr
        for tag, value in (("w:tblW", str(sum(widths))), ("w:tblInd", "120")):
            node = properties.find(qn(tag))
            if node is None:
                node = OxmlElement(tag)
                properties.append(node)
            node.set(qn("w:w"), value)
            node.set(qn("w:type"), "dxa")
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                tc_pr = cell._tc.get_or_add_tcPr()
                width_node = tc_pr.find(qn("w:tcW"))
                if width_node is None:
                    width_node = OxmlElement("w:tcW")
                    tc_pr.append(width_node)
                width_node.set(qn("w:w"), str(widths[index]))
                width_node.set(qn("w:type"), "dxa")
                margins = tc_pr.find(qn("w:tcMar"))
                if margins is None:
                    margins = OxmlElement("w:tcMar")
                    tc_pr.append(margins)
                for side, amount in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                    edge = margins.find(qn(f"w:{side}"))
                    if edge is None:
                        edge = OxmlElement(f"w:{side}")
                        margins.append(edge)
                    edge.set(qn("w:w"), str(amount))
                    edge.set(qn("w:type"), "dxa")

    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    normal.font.name, normal.font.size = "Calibri", Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, _ACCENT, 16, 8),
        ("Heading 2", 13, _ACCENT, 12, 6),
        ("Heading 3", 12, _INK, 8, 4),
    ):
        style = document.styles[name]
        style.font.name, style.font.size = "Calibri", Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    core = document.core_properties
    core.title = "Campaign Evidence Report"
    core.author = "BashGym Campaign Controller"
    core.created = core.modified = _FIXED_TIME
    header = section.header.paragraphs[0]
    set_font(header.add_run("BASHGYM | CAMPAIGN EVIDENCE"), 8.5, _MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("Reconciled evidence package"), 8.5, _MUTED)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run("CAMPAIGN EVIDENCE REPORT"), 24, _INK, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_font(subtitle.add_run(str(snapshot["campaign"].get("objective", "Experiment campaign"))), 13, _MUTED)
    metadata = (
        ("Campaign", snapshot["campaign"].get("campaign_id", "unknown")),
        ("Status", snapshot["campaign"].get("status", "unknown")),
        ("Champion", snapshot["campaign"].get("champion_ref") or "unchanged / not recorded"),
        ("Evidence digest", source_digest),
    )
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        set_font(paragraph.add_run(f"{label}: "), 10.5, _INK, bold=True)
        set_font(paragraph.add_run(str(value)), 10.5, _INK)

    document.add_heading("Model-quality findings", level=1)
    document.add_paragraph(_quality_text(snapshot))
    document.add_heading("Training evidence", level=1)
    picture = document.add_picture(str(loss_png), width=Inches(6.5))
    picture._inline.docPr.set(
        "descr",
        "Training loss by persisted step; dashed series identify smoke engineering evidence.",
    )
    picture._inline.docPr.set("title", "Campaign training loss")
    caption = document.add_paragraph("Figure 1. Persisted loss series. Dashed lines are smoke engineering evidence.")
    caption.paragraph_format.space_before, caption.paragraph_format.space_after = Pt(4), Pt(4)
    set_font(caption.runs[0], 9, _MUTED)

    document.add_page_break()
    document.add_heading("Attempts", level=1)
    table = document.add_table(rows=1, cols=4)
    headers = ("Attempt", "Stage", "Status", "Candidate")
    for cell, label in zip(table.rows[0].cells, headers, strict=True):
        shade(cell, _TABLE_FILL)
        set_font(cell.paragraphs[0].add_run(label), 9, _INK, bold=True)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_marker = OxmlElement("w:tblHeader")
    header_marker.set(qn("w:val"), "1")
    header_properties.append(header_marker)
    for item in snapshot["attempts"]:
        values = (
            item.get("attempt_id", ""), item.get("stage", ""), item.get("status", ""),
            str(item.get("candidate_digest", ""))[:12],
        )
        cells = table.add_row().cells
        for cell, value in zip(cells, values, strict=True):
            set_font(cell.paragraphs[0].add_run(str(value)), 8.5, _INK)
    set_table_geometry(table, (2250, 2450, 1800, 2860))

    document.add_heading("Evidence and flags", level=1)
    document.add_paragraph(
        f"{len(snapshot['artifacts'])} sealed artifact records and "
        f"{len(snapshot['comparisons'])} comparison records are included in the evidence package."
    )
    flags = snapshot["flags"] or ["No implementation flags were recorded for this export."]
    for flag in flags:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.space_after = Pt(8)
        set_font(paragraph.add_run(str(flag)), 11, _INK)
    document.add_heading("Reconciliation", level=1)
    document.add_paragraph(
        "Every table and chart is derived from campaign_evidence.json. "
        "export_manifest.json records the SHA-256 of every generated file."
    )
    document.save(path)
    _normalize_docx(path)


def write_campaign_pdf(
    snapshot: dict[str, Any], source_digest: str, loss_png: Path, path: Path
) -> None:
    """Create a deterministic PDF companion to the editable report."""

    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.pdfgen.canvas import Canvas
        from reportlab.platypus import (
            Image,
            PageBreak,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError as exc:  # pragma: no cover - installation-dependent
        raise CampaignReportingUnavailableError("Install bashgym[reports] for PDF exports") from exc

    class InvariantCanvas(Canvas):
        def __init__(self, *args, **kwargs):
            kwargs["invariant"] = 1
            super().__init__(*args, **kwargs)

    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "CampaignBody", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5,
        leading=14, spaceAfter=6, textColor=colors.HexColor(f"#{_INK}"), alignment=TA_LEFT,
    )
    heading = ParagraphStyle(
        "CampaignHeading", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=15,
        leading=18, spaceBefore=14, spaceAfter=7, textColor=colors.HexColor(f"#{_ACCENT}"),
    )
    title = ParagraphStyle(
        "CampaignTitle", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=23,
        leading=27, spaceAfter=5, textColor=colors.HexColor(f"#{_INK}"), alignment=TA_LEFT,
    )
    subtitle = ParagraphStyle(
        "CampaignSubtitle", parent=body, fontSize=12.5, leading=16, spaceAfter=14,
        textColor=colors.HexColor(f"#{_MUTED}"),
    )
    document = SimpleDocTemplate(
        str(path), pagesize=letter, leftMargin=inch, rightMargin=inch,
        topMargin=inch, bottomMargin=inch, title="Campaign Evidence Report",
        author="BashGym Campaign Controller",
    )
    campaign = snapshot["campaign"]
    story = [
        Paragraph("CAMPAIGN EVIDENCE REPORT", title),
        Paragraph(str(campaign.get("objective", "Experiment campaign")), subtitle),
    ]
    for label, value in (
        ("Campaign", campaign.get("campaign_id", "unknown")),
        ("Status", campaign.get("status", "unknown")),
        ("Champion", campaign.get("champion_ref") or "unchanged / not recorded"),
        ("Evidence digest", source_digest),
    ):
        story.append(Paragraph(f"<b>{label}:</b> {value}", body))
    story.extend(
        [
            Paragraph("Model-quality findings", heading),
            Paragraph(_quality_text(snapshot), body),
            Paragraph("Training evidence", heading),
            Image(str(loss_png), width=6.5 * inch, height=3.25 * inch),
            Paragraph("Figure 1. Persisted loss series. Dashed lines are smoke engineering evidence.", body),
            PageBreak(),
            Paragraph("Attempts", heading),
        ]
    )
    rows = [["Attempt", "Stage", "Status", "Candidate"]]
    rows.extend(
        [
            str(item.get("attempt_id", "")), str(item.get("stage", "")),
            str(item.get("status", "")), str(item.get("candidate_digest", ""))[:12],
        ]
        for item in snapshot["attempts"]
    )
    table = Table(rows, colWidths=[1.55 * inch, 1.75 * inch, 1.15 * inch, 2.05 * inch], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{_TABLE_FILL}")),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor(f"#{_INK}")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8B0A8")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend([table, Spacer(1, 8), Paragraph("Evidence and flags", heading)])
    story.append(
        Paragraph(
            f"{len(snapshot['artifacts'])} sealed artifact records and "
            f"{len(snapshot['comparisons'])} comparison records are included.", body,
        )
    )
    flags = snapshot["flags"] or ["No implementation flags were recorded for this export."]
    for flag in flags:
        story.append(Paragraph(f"&#8226; {flag}", body))
    story.extend(
        [
            Paragraph("Reconciliation", heading),
            Paragraph(
                "Every table and chart is derived from campaign_evidence.json. "
                "export_manifest.json records the SHA-256 of every generated file.", body,
            ),
        ]
    )
    document.build(story, canvasmaker=InvariantCanvas)


__all__ = [
    "CampaignReportingUnavailableError",
    "write_campaign_docx",
    "write_campaign_pdf",
    "write_loss_png",
]
